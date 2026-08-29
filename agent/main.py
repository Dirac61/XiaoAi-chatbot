from fastapi import FastAPI
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from openai import AsyncOpenAI
from dotenv import load_dotenv
from contextlib import asynccontextmanager
import os
import httpx
import logging
import asyncio
import json
import time
import uuid
import secrets
import contextvars
from enum import Enum
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Tuple, AsyncIterator, AsyncGenerator, Callable, Awaitable

# ---- correlation_id：每条请求分配一个短ID，所有日志自动带 [req-XXXX] 前缀 ----
# 解决：多并发对话时日志糊在一起，分不清哪条请求的流水
_REQ_ID_CTX: contextvars.ContextVar[str] = contextvars.ContextVar("xiaoai_req_id", default="")


class _CorrelationFilter(logging.Filter):
    """
    把 _REQ_ID_CTX 当前值注入到日志 record 的 req_id 字段。
    空字符串（服务启动期/非请求上下文）就显示为「boot」，避免启动日志出现空括号。
    """

    def filter(self, record: logging.LogRecord) -> bool:
        ctx_id = _REQ_ID_CTX.get() or ""
        record.req_id = f"[req-{ctx_id}]" if ctx_id else "[boot]"
        return True


# 初始化 root logger 的基础配置（只做一次，避免重复 handler）
if not logging.getLogger().handlers:
    _stream_handler = logging.StreamHandler()
    # 日志格式（更紧凑 + 带 req_id）：时间 | 级别 | req_id | logger名 | 内容
    _stream_handler.setFormatter(logging.Formatter(
        "%(asctime)s | %(levelname)-7s | %(req_id)s | %(name)s | %(message)s",
        datefmt="%H:%M:%S",
    ))
    _stream_handler.addFilter(_CorrelationFilter())
    logging.getLogger().addHandler(_stream_handler)
    logging.getLogger().setLevel(logging.INFO)

# MemoryService 之前开 DEBUG 会刷屏（每条向量命中、BM25 分数、Redis 去重细节），
# 按用户要求"日志太乱"默认调回 INFO；需要排查时再手动改为 DEBUG。
logging.getLogger("XiaoAi Memory Service").setLevel(logging.INFO)

# 把 correlation filter 挂到几个关键 logger 上，防止部分平台 logger 不继承 root handler 的 filter
for _logger_name in ("XiaoAi Agent", "XiaoAi Memory Service", "uvicorn", "uvicorn.access"):
    _lg = logging.getLogger(_logger_name)
    for _h in _lg.handlers:
        if not any(isinstance(_f, _CorrelationFilter) for _f in _h.filters):
            _h.addFilter(_CorrelationFilter())

logger = logging.getLogger("XiaoAi Agent")


def new_req_id() -> str:
    """生成一个短且够用的请求ID（6字节hex=12字符，碰撞极低）。"""
    return secrets.token_hex(3)


def bind_req_id(req_id: Optional[str] = None) -> str:
    """把请求ID绑定到当前上下文，返回被绑定的ID。"""
    _id = req_id if req_id else new_req_id()
    _REQ_ID_CTX.set(_id)
    return _id

load_dotenv()

from services.memory_service import memory_service
from services.search_service import search_service

# ========== 【硬编码迁移】集中配置导入 ==========
# 所有 prompt 大文本统一放到 config.prompts.*，避免 main.py 里散落 2000 行字符串
from config.prompts.persona import (
    PERSONA_SYSTEM_PROMPT,
    PERSONA_MEMORY_SECTION_TITLE,
    PERSONA_SEARCH_SECTION_TITLE,
)
from config.prompts.ocr import (
    OCR_SINGLE_IMAGE_SYSTEM_PROMPT, OCR_SINGLE_IMAGE_USER_TEXT,
    OCR_MULTI_IMAGE_SYSTEM_PROMPT, OCR_MULTI_IMAGE_USER_TEXT,
    OCR_DISABLE_THINKING,
)
from config.prompts.orchestrator_fast import (
    FAST_ORCH_SYSTEM_PROMPT, build_fast_orch_user_prompt,
)
from config.prompts.orchestrator_expert import (
    build_expert_orch_system_prompt, build_expert_orch_user_prompt,
    EXPERT_ORCH_USER_RETRY_APPEND,
)
from config.prompts.deep_thinking import (
    DEEP_THINKING_SYSTEM_PROMPT,
    DEEP_THINKING_START_MESSAGE,
    DEEP_THINKING_REASONING_EFFORT,
)
from config.prompts.reply_directly import (
    REPLY_DIRECTLY_SYSTEM_PROMPT,
    build_reply_directly_user_prompt,
)
from config.prompts.summarize import (
    SUMMARIZE_SYSTEM_PROMPT,
    build_summarize_user_prompt,
    SUMMARIZE_DISABLE_THINKING,
)
from config.prompts.context import (
    build_global_context_text,
    build_thinking_history_text,
    format_recent_history_for_llm,
    USER_PREFIX_UPLOAD_SINGLE_IMAGE,
    USER_PREFIX_UPLOAD_SINGLE_IMAGE_WITH_EXTRACTED,
    USER_PREFIX_UPLOAD_SINGLE_FILE,
    USER_PREFIX_UPLOAD_MULTI_FILES,
    HISTORY_IMAGE_PREFIX_SIMPLE,
    HISTORY_IMAGE_PREFIX_WITH_EXTRACTED,
    HISTORY_FILE_PREFIX_SIMPLE,
    HISTORY_FILE_PREFIX_WITH_EXTRACTED,
    FILE_TEXT_BLOCK_PREFIX,
)

# 行为参数（温度/截断/权重/重试/超时/top_k）统一放到 config.behavior
from config.behavior import (
    # 文本截断
    MEDIA_EXTRACTED_MAX_LEN,
    FAST_ORCH_CONTEXT_MAX_LEN,
    HISTORY_SNIPPET_MAX_LEN_EACH,
    LOG_ANALYSIS_TRUNCATE_LEN,
    LOG_USER_MESSAGE_TRUNCATE_LEN,
    # 各模型温度 / max_tokens / 重试
    OCR_TEMPERATURE,
    FAST_ORCH_TEMPERATURE, FAST_ORCH_MAX_RETRIES,
    EXPERT_ORCH_TEMPERATURE, EXPERT_ORCH_MAX_TOKENS, EXPERT_ORCH_MAX_RETRIES,
    DEEP_THINKING_TEMPERATURE,
    REPLY_DIRECTLY_TEMPERATURE, REPLY_DIRECTLY_MAX_TOKENS,
    SUMMARIZE_TEMPERATURE,
    MEMORY_SEARCH_DEFAULT_TOP_K,
    # 后端 / HTTP 超时
    BACKEND_INTERNAL_API_TIMEOUT, BACKEND_INTERNAL_API_CONNECT_TIMEOUT,
    CLIENT_TIMEOUT_READ_GENERAL, CLIENT_TIMEOUT_CONNECT_DEFAULT,
    CLIENT_TIMEOUT_READ_ORCHESTRATION, CLIENT_TIMEOUT_READ_DEEP_THINKING,
    # 文件扩展名白名单 / URL 前缀判断
    PLAIN_TEXT_EXTENSIONS, SUPPORTED_FILE_EXTENSIONS,
    URL_PREFIX_DATA_IMAGE, URL_PREFIX_HTTP,
    # 专家模式 dataclass 兜底默认值
    EXPERT_STATE_DEFAULT_MAX_ITERATIONS, EXPERT_STATE_DEFAULT_HISTORY,
)

# 工具协议常量（动作枚举、工具名、SSE 摘要模板、后端接口后缀、mediaType）
from config.tools import (
    EXPERT_ACTION_COLLECT, EXPERT_ACTION_DEEP, EXPERT_ACTION_REPLY,
    EXPERT_FINAL_PATH_DEEP, EXPERT_FINAL_PATH_REPLY, FAST_MODE_FINAL_PATH,
    TOOL_NAME_WEB_SEARCH, TOOL_NAME_MEMORY_SEARCH,
    TOOL_DESC_WEB_SEARCH, TOOL_DESC_MEMORY_SEARCH,
    SUMMARY_TEMPLATE_WEB_SEARCH, SUMMARY_TEMPLATE_MEMORY_SEARCH,
    SUMMARY_TEMPLATE_WEB_SEARCH_FAIL, SUMMARY_TEMPLATE_MEMORY_SEARCH_FAIL,
    ORCH_PHASE_PLANNING, ORCH_PHASE_THINKING,
    BACKEND_ENDPOINT_UPDATE_SEARCH_RESULTS,
    BACKEND_ENDPOINT_UPDATE_MESSAGE_CONTENT,
    BACKEND_ENDPOINT_UPDATE_EXPERT_TRACE_LEGACY,
    BACKEND_ENDPOINT_MEMORY_DELETE,
    MEDIA_TYPE_IMAGE, MEDIA_TYPE_FILE,
    default_tool_description_lines,
)


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("初始化MemoryService...")
    memory_service.init()
    yield


app = FastAPI(title="XiaoAi Agent", version="1.0.0", lifespan=lifespan)

API_KEY = os.getenv("API_KEY", "")
MODEL = os.getenv("MODEL", "")
API_BASE = os.getenv("API_BASE", "https://api.openai.com/v1")

MULTIMODAL_MODEL = os.getenv("MULTIMODAL_MODEL", MODEL)
MULTIMODAL_API_KEY = os.getenv("MULTIMODAL_API_KEY", API_KEY)
MULTIMODAL_API_BASE = os.getenv("MULTIMODAL_API_BASE", API_BASE)
USE_MULTIMODAL = os.getenv("USE_MULTIMODAL", "false").lower() == "true"

OCR_MODEL = os.getenv("OCR_MODEL") or MULTIMODAL_MODEL
OCR_API_KEY = os.getenv("OCR_API_KEY") or MULTIMODAL_API_KEY
OCR_API_BASE = os.getenv("OCR_API_BASE") or MULTIMODAL_API_BASE

ORCHESTRATION_MODEL = os.getenv("ORCHESTRATION_MODEL") or MODEL
ORCHESTRATION_API_KEY = os.getenv("ORCHESTRATION_API_KEY") or API_KEY
ORCHESTRATION_API_BASE = os.getenv("ORCHESTRATION_API_BASE") or API_BASE

logger.info(f"文本模型 - API_KEY: {'是' if API_KEY else '否'}, 模型: {MODEL}, 地址: {API_BASE}")
logger.info(f"多模态模型 - API_KEY: {'是' if MULTIMODAL_API_KEY else '否'}, 模型: {MULTIMODAL_MODEL}, 地址: {MULTIMODAL_API_BASE}, 启用: {'是' if USE_MULTIMODAL else '否'}")
logger.info(f"OCR模型 - API_KEY: {'是' if OCR_API_KEY else '否'}, 模型: {OCR_MODEL}, 地址: {OCR_API_BASE}")
logger.info(f"编排器模型 - API_KEY: {'是' if ORCHESTRATION_API_KEY else '否'}, 模型: {ORCHESTRATION_MODEL}, 地址: {ORCHESTRATION_API_BASE}")

# === 专家模式配置 ===
EXPERT_MODE_ENABLED = os.getenv("EXPERT_MODE_ENABLED", "false").lower() == "true"
EXPERT_MAX_ITERATIONS = int(os.getenv("EXPERT_MAX_ITERATIONS", "3"))

EXPERT_ORCHESTRATION_MODEL = os.getenv("EXPERT_ORCHESTRATION_MODEL", "")
EXPERT_ORCHESTRATION_API_KEY = os.getenv("EXPERT_ORCHESTRATION_API_KEY", "")
EXPERT_ORCHESTRATION_API_BASE = os.getenv("EXPERT_ORCHESTRATION_API_BASE", "")

EXPERT_DEEP_THINKING_MODEL = os.getenv("EXPERT_DEEP_THINKING_MODEL", "")
EXPERT_DEEP_THINKING_API_KEY = os.getenv("EXPERT_DEEP_THINKING_API_KEY", "")
EXPERT_DEEP_THINKING_API_BASE = os.getenv("EXPERT_DEEP_THINKING_API_BASE", "")

logger.info(f"专家模式 - 启用: {'是' if EXPERT_MODE_ENABLED else '否'}")
logger.info(f"专家模式编排器 - API_KEY: {'是' if EXPERT_ORCHESTRATION_API_KEY else '否'}, 模型: {EXPERT_ORCHESTRATION_MODEL}, 地址: {EXPERT_ORCHESTRATION_API_BASE}")
logger.info(f"深度思考模型 - API_KEY: {'是' if EXPERT_DEEP_THINKING_API_KEY else '否'}, 模型: {EXPERT_DEEP_THINKING_MODEL}, 地址: {EXPERT_DEEP_THINKING_API_BASE}")
logger.info(f"专家模式最大迭代次数: {EXPERT_MAX_ITERATIONS}")
logger.info(f"专家模式 - 编排器直接生成回复，不使用主模型")

# === 后端内部回写配置 ===
# 设计意图：三个内部接口（更新消息内容/搜索结果/expertTrace）都需要：
#  1) 后端服务地址（默认开发时为 http://localhost:8080，可通过 env 覆盖为内网部署域名）
#  2) X-Internal-Secret 鉴权密钥（project memory 约定：内部接口必须通过 X-Internal-Secret 鉴权）
# 为避免各回写函数中散落 os.getenv 出现"新函数漏写导致 NameError"（本次报错即此因），
# 统一在顶部一次性读取为全局常量，后续回写函数直接引用。
BACKEND_BASE_URL: str = os.getenv("BACKEND_URL", "http://localhost:8080").rstrip("/")
INTERNAL_SECRET: str = os.getenv("INTERNAL_SECRET", "")
# 校验日志：密钥长度只打印是否为空，不打印内容（安全）
logger.info(f"后端服务地址: {BACKEND_BASE_URL}")
logger.info(f"内部接口密钥: {'已配置' if INTERNAL_SECRET else '⚠未配置(内部接口会被后端拦截)'}")

# 主聊天模型 client
# timeout: 读取超时 / connect 超时统一用 config.behavior 的 CLIENT_TIMEOUT_*
# max_retries=0: 禁用 AsyncOpenAI 内置重试，避免单次超时被放大为 3 倍时长，失败由业务层处理
client = AsyncOpenAI(
    api_key=API_KEY,
    base_url=API_BASE,
    timeout=httpx.Timeout(CLIENT_TIMEOUT_READ_GENERAL, connect=CLIENT_TIMEOUT_CONNECT_DEFAULT),
    max_retries=0
)

# 多模态模型 client（配置策略同主 client）
multimodal_client = None
if MULTIMODAL_API_KEY and MULTIMODAL_API_BASE:
    multimodal_client = AsyncOpenAI(
        api_key=MULTIMODAL_API_KEY,
        base_url=MULTIMODAL_API_BASE,
        timeout=httpx.Timeout(CLIENT_TIMEOUT_READ_GENERAL, connect=CLIENT_TIMEOUT_CONNECT_DEFAULT),
        max_retries=0
    )

# OCR 模型 client（配置策略同主 client）
ocr_client = None
if OCR_API_KEY and OCR_API_BASE:
    ocr_client = AsyncOpenAI(
        api_key=OCR_API_KEY,
        base_url=OCR_API_BASE,
        timeout=httpx.Timeout(CLIENT_TIMEOUT_READ_GENERAL, connect=CLIENT_TIMEOUT_CONNECT_DEFAULT),
        max_retries=0
    )

# 编排器小模型 client
# timeout: 读取走 CLIENT_TIMEOUT_READ_ORCHESTRATION（编排器只输出 JSON，响应应较快）
# max_retries=0: 禁用重试，失败时直接走 fallback（need_search=False），不拖累主流程
orchestration_client = None
if ORCHESTRATION_API_KEY and ORCHESTRATION_API_BASE:
    orchestration_client = AsyncOpenAI(
        api_key=ORCHESTRATION_API_KEY,
        base_url=ORCHESTRATION_API_BASE,
        timeout=httpx.Timeout(CLIENT_TIMEOUT_READ_ORCHESTRATION, connect=CLIENT_TIMEOUT_CONNECT_DEFAULT),
        max_retries=0
    )

# === 专家模式专用 client ===
# 专家模式编排器 client（多模态模型，支持图片理解）
expert_orchestration_client = None
if EXPERT_ORCHESTRATION_API_KEY and EXPERT_ORCHESTRATION_API_BASE:
    expert_orchestration_client = AsyncOpenAI(
        api_key=EXPERT_ORCHESTRATION_API_KEY,
        base_url=EXPERT_ORCHESTRATION_API_BASE,
        timeout=httpx.Timeout(CLIENT_TIMEOUT_READ_GENERAL, connect=CLIENT_TIMEOUT_CONNECT_DEFAULT),
        max_retries=0
    )

# 深度思考模型 client（多模态模型，接收图片URL进行分析，生成慢所以读取超时更宽）
deep_thinking_client = None
if EXPERT_DEEP_THINKING_API_KEY and EXPERT_DEEP_THINKING_API_BASE:
    deep_thinking_client = AsyncOpenAI(
        api_key=EXPERT_DEEP_THINKING_API_KEY,
        base_url=EXPERT_DEEP_THINKING_API_BASE,
        timeout=httpx.Timeout(CLIENT_TIMEOUT_READ_DEEP_THINKING, connect=CLIENT_TIMEOUT_CONNECT_DEFAULT),
        max_retries=0
    )

# 专家模式主模型 client - 已移除，由编排器直接生成回复


# ==============================================================================
# 专家模式：状态机核心（DECIDING / EXECUTING / FINAL / DONE）
# ==============================================================================

class ExpertPhase(str, Enum):
    """专家模式状态机的 4 个核心阶段"""
    DECIDING = "deciding"     # 编排器输出单步计划（action + step + analysis）
    EXECUTING = "executing"    # 并行执行 step.tools[]
    FINAL = "final"           # 二选一流式正文：深度思考 或 回复模型
    DONE = "done"             # 结束


# ---- 编排器允许的 action 枚举 ----
# 【硬编码移除】action 常量不再在 main.py 里重复声明；直接用 config.tools 里导入的唯一真源：
#   EXPERT_ACTION_COLLECT / EXPERT_ACTION_DEEP / EXPERT_ACTION_REPLY
# 这样改工具动作名称或新增动作时，只改 config.tools 一处，main.py / 前端协议 / prompts 自动同步。


@dataclass
class ToolCall:
    """一次具体的工具调用请求（编排器 step.tools[] 中的一个元素）"""
    tool: str
    params: Dict[str, Any] = field(default_factory=dict)


@dataclass
class PlanStep:
    """编排器输出的单步计划对象（一次一步，不再是数组）"""
    purpose: str
    tools: List[ToolCall] = field(default_factory=list)


@dataclass
class OrchResult:
    """编排器一次输出的结构化结果"""
    action: str                       # collect_tools / deep_thinking / reply_directly
    step: Optional[PlanStep]          # action=collect_tools 时必填；其他为 None
    analysis: str                     # 一句话/多句话分析，下轮编排器用来了解"之前怎么想的"
    raw_output: str = ""              # 原始输出字符串，仅用于日志排错


@dataclass
class ToolExecutionRecord:
    """单个工具的执行结果（存入持久化 expertTrace.history[].tools[]，不做摘要）"""
    tool: str                        # 工具名
    params: Dict[str, Any]           # 调用入参（用户要求"不进行摘要，本身就简洁"）
    success: bool                    # 是否成功
    durationMs: int                  # 耗时 ms
    resultCount: Optional[int] = None  # 搜索条数 / 记忆条数 / None
    error: Optional[str] = None      # 失败原因（非空说明失败）
    # 为了"工具调用结果也保存不摘要"：额外保留对最终回复有价值的精简原始结果片段
    #  web_search 时：过滤后的 [{title,url}] 列表
    #  memory_search 时：[content 截断前100字] 列表
    rawResult: Any = None            # 精简原始结果（不是摘要，是实际可展示的原始精简数据）


@dataclass
class OrchHistoryRecord:
    """一次 DECIDING + 对应 EXECUTING（或 FINAL）的持久化记录"""
    iteration: int
    action: str
    purpose: Optional[str]
    analysis: str
    tools: List[ToolExecutionRecord] = field(default_factory=list)


@dataclass
class ExpertState:
    """专家模式状态机的唯一全局数据源（替代原函数内散落的局部变量）。

    ⚠ Python dataclass 硬性规则：**无默认值字段必须全部放在有默认值字段之前**。
    所以将"必填输入（message / history / user_id ... / request_start_time）"统一前置，
    "双 UUID / 状态机 / 累积产物"这些带默认值的字段统一次后，避免出现：
      TypeError: non-default argument 'request_start_time' follows default argument 'message_uuid'
    """
    # ===== 第一组：必填输入（无默认值）=====
    message: str
    message_type: str
    media_urls: List[str]
    history: List[dict]
    user_id: Optional[int]
    session_id: Optional[str]
    request_start_time: float

    # ===== 第二组：可选/衍生输入（有默认值）=====
    # 双 UUID（与后端 ChatController 对齐，修复"回写字段覆盖到用户消息"Bug）：
    #  - user_message_uuid     : 本次用户消息的 UUID（role=user）→ 用于更新正文/提取文本（update_backend_message_content）
    #  - message_uuid          : 本次 assistant 回复的 UUID（role=assistant）→ 用于写入 searchResults / expertTrace
    #    （保留原字段名 message_uuid，是因为 expertTrace 回写处 state.message_uuid 被大量使用；语义调整为"assistant uuid"）
    user_message_uuid: Optional[str] = None
    message_uuid: Optional[str] = None

    # ===== 第三组：预处理产物（有默认值）=====
    extracted_text: str = ""
    memories: List[dict] = field(default_factory=list)

    # ===== 第四组：状态机核心（有默认值）=====
    phase: ExpertPhase = ExpertPhase.DECIDING
    iteration: int = 0
    # 【硬编码移除】原默认值 2 散落 dataclass 内部；统一走 config.behavior.EXPERT_STATE_DEFAULT_MAX_ITERATIONS
    max_iterations: int = EXPERT_STATE_DEFAULT_MAX_ITERATIONS  # 构造时仍会被 EXPERT_MAX_ITERATIONS env 覆盖
    deep_thinking_called: bool = False

    # 当前步（由 DECIDING 产生，EXECUTING 消费）
    pending_step: Optional[PlanStep] = None
    pending_orch: Optional[OrchResult] = None

    # ===== 第五组：累积产物（有默认值）=====
    search_results_flattened: List[dict] = field(default_factory=list)
    all_search_context_parts: List[str] = field(default_factory=list)
    memory_context_parts: List[str] = field(default_factory=list)

    # 持久化数据：编排历史（analysis）+ 工具执行历史（原始结果不摘要）
    orch_history: List[OrchHistoryRecord] = field(default_factory=list)

    # FINAL 阶段流式结果（正文 + 深度思考reasoning）
    final_path: Optional[str] = None   # deep_thinking / reply_directly
    assistant_response: str = ""
    # 深度思考推理链全文（用户要求持久化到 expertTrace，与工具调用结果同等对待）
    deep_thinking_reasoning: str = ""


# ==============================================================================
# 专家模式：工具注册表（模块化扩展工具，不改状态机主循环）
# 工具签名：fn(params, state) -> (ToolExecutionRecord, stream_chunks list[str])
#   stream_chunks：要立刻向前端发送的 JSON 行（如 search_start / search_summary）
# ==============================================================================

ExpertToolFn = Callable[[Dict[str, Any], "ExpertState"], Awaitable[Tuple[ToolExecutionRecord, List[str]]]]
EXPERT_TOOL_REGISTRY: Dict[str, Tuple[ExpertToolFn, str]] = {}   # tool_name -> (fn, description)


def _register_tool(name: str, description: str, fn: ExpertToolFn):
    """注册一个专家模式工具（模块化扩展用）。工具注册只在启动时执行一次，降为 DEBUG 避免刷屏。"""
    EXPERT_TOOL_REGISTRY[name] = (fn, description)
    logger.debug(f"[专家模式][工具注册] {name}: {description}")


async def _tool_web_search(params: Dict[str, Any], state: ExpertState) -> Tuple[ToolExecutionRecord, List[str]]:
    """工具：联网搜索；返回执行记录 + 前端要立刻发送的流式消息行。
    【硬编码移除】
      - 工具名字符串 → TOOL_NAME_WEB_SEARCH
      - SSE 摘要模板 → SUMMARY_TEMPLATE_WEB_SEARCH / SUMMARY_TEMPLATE_WEB_SEARCH_FAIL
    """
    start = time.time()
    keywords: List[str] = params.get("keywords") or []
    stream_chunks: List[str] = []
    error: Optional[str] = None
    result_count = 0
    raw_result: List[dict] = []  # 过滤后的 [{title,url}]

    if not keywords:
        error = "keywords 为空"
        logger.warning("[专家模式][工具:web_search] keywords 为空，跳过")
    else:
        try:
            # 立即向前端发送 tool_call_start 状态（统一工具调用协议）
            stream_chunks.append(json.dumps({
                "type": "tool_call_start",
                "data": {"tool": TOOL_NAME_WEB_SEARCH, "params": {"keywords": keywords}},
            }, ensure_ascii=False))
            # 单工具启动降到 DEBUG：每次专家迭代如果有工具调用就会打 2 条（开始+完成），
            # 汇总交给 _expert_execute_step 的 [EXECUTING] 完成行，避免日志与 summary 行重复
            logger.debug("[专家模式][工具:web_search] 开始搜索 keywords=%s", keywords)

            search_start = time.time()
            results = await search_service.web_search(keywords)
            context_text = await search_service.get_search_context(keywords)
            duration_s = time.time() - search_start

            # 过滤结果（只保留 title/url，项目既有约束）
            raw_result = [
                {"title": r.get("title", ""), "url": r.get("url", "")}
                for r in (results or [])
                if r.get("url")
            ]
            result_count = len(raw_result)

            # 累积到全局 state（供 FINAL 阶段 deep_thinking / reply_model 使用）
            state.search_results_flattened.extend(raw_result)
            if context_text:
                state.all_search_context_parts.append(context_text)

            # web_search 完成降 DEBUG：_expert_execute_step 会在一次迭代的多工具并行完成后
            # 打 1 条汇总行（包含所有工具名+各自结果计数），单工具不必再 INFO 刷屏
            logger.debug("[专家模式][工具:web_search] 完成 count=%s, 耗时=%.2fs", result_count, duration_s)

            # tool_call_result 流式块（统一工具调用协议，前端据此渲染工具摘要行）
            stream_chunks.append(json.dumps({
                "type": "tool_call_result",
                "data": {
                    "tool": TOOL_NAME_WEB_SEARCH,
                    "success": True,
                    "durationMs": int(duration_s * 1000),
                    "resultCount": result_count,
                    "summary": SUMMARY_TEMPLATE_WEB_SEARCH.format(n=result_count),
                    "results": raw_result,
                },
            }, ensure_ascii=False))
        except Exception as e:
            error = f"{type(e).__name__}: {str(e)}"
            logger.error(f"[专家模式][工具:web_search] 失败: {error}")
            # 异常时也发 tool_call_result，前端能感知失败
            stream_chunks.append(json.dumps({
                "type": "tool_call_result",
                "data": {
                    "tool": TOOL_NAME_WEB_SEARCH,
                    "success": False,
                    "durationMs": 0,
                    "resultCount": 0,
                    "summary": SUMMARY_TEMPLATE_WEB_SEARCH_FAIL,
                    "error": error,
                },
            }, ensure_ascii=False))

    duration_ms = int((time.time() - start) * 1000)
    record = ToolExecutionRecord(
        tool=TOOL_NAME_WEB_SEARCH,
        params={"keywords": keywords},
        success=(error is None),
        durationMs=duration_ms,
        resultCount=result_count,
        error=error,
        rawResult=raw_result,
    )
    return record, stream_chunks


async def _tool_memory_search(params: Dict[str, Any], state: ExpertState) -> Tuple[ToolExecutionRecord, List[str]]:
    """工具：长期记忆检索（静默，不向前端发流式消息）。
    【硬编码移除】
      - 工具名字符串 → TOOL_NAME_MEMORY_SEARCH
      - top_k=5 → MEMORY_SEARCH_DEFAULT_TOP_K（与 fast 路径一致单源）
      - SSE 摘要模板 → SUMMARY_TEMPLATE_MEMORY_SEARCH / SUMMARY_TEMPLATE_MEMORY_SEARCH_FAIL
    """
    start = time.time()
    query: str = params.get("query") or state.message
    stream_chunks: List[str] = []
    error: Optional[str] = None
    result_count = 0
    raw_result: List[str] = []

    # ===== 新增：tool_call_start 通知（统一工具调用协议，前端显示工具正在运行）=====
    stream_chunks.append(json.dumps({
        "type": "tool_call_start",
        "data": {"tool": TOOL_NAME_MEMORY_SEARCH, "params": {"query": query[:50]}},
    }, ensure_ascii=False))

    try:
        # memory_search 启动：静默工具（不向前端发 SSE），每次只打 DEBUG 避免 INFO 堆
        logger.debug("[专家模式][工具:memory_search] query=%s...", (query[:50] if query else ""))
        items = []
        if state.session_id:
            # 【硬编码移除】top_k=5 → MEMORY_SEARCH_DEFAULT_TOP_K（快/专家模式一致）
            items = await memory_service.search_memories(query, str(state.session_id), top_k=MEMORY_SEARCH_DEFAULT_TOP_K)
        result_count = len(items)
        raw_result = [(m.get("content") or "")[:200] for m in items]

        # 累积到 state
        if items:
            parts = [f"- {m.get('content','')[:300]}" for m in items]
            state.memory_context_parts.extend(parts)
    except Exception as e:
        error = f"{type(e).__name__}: {str(e)}"
        logger.error(f"[专家模式][工具:memory_search] 失败: {error}")

    duration_ms = int((time.time() - start) * 1000)

    # ===== 新增：tool_call_result 通知（前端据此渲染"已读取 N 个文件"摘要）=====
    stream_chunks.append(json.dumps({
        "type": "tool_call_result",
        "data": {
            "tool": TOOL_NAME_MEMORY_SEARCH,
            "success": (error is None),
            "durationMs": duration_ms,
            "resultCount": result_count,
            "summary": SUMMARY_TEMPLATE_MEMORY_SEARCH.format(n=result_count),
            "error": error,
        },
    }, ensure_ascii=False))

    record = ToolExecutionRecord(
        tool=TOOL_NAME_MEMORY_SEARCH,
        params={"query": query},
        success=(error is None),
        durationMs=duration_ms,
        resultCount=result_count,
        error=error,
        rawResult=raw_result,
    )
    return record, stream_chunks


# 【硬编码移除】工具注册名/描述改用 config.tools 常量（单源；后续改工具名只改一处）
_register_tool(TOOL_NAME_WEB_SEARCH, TOOL_DESC_WEB_SEARCH, _tool_web_search)
_register_tool(TOOL_NAME_MEMORY_SEARCH, TOOL_DESC_MEMORY_SEARCH, _tool_memory_search)


# ==============================================================================
# 工具函数：生成 JSON 行 / 构建思考历史 / 构建全局上下文 / 关思考统一开关
# ==============================================================================

def _json_line(obj: Dict[str, Any]) -> str:
    """把对象序列化成流式响应的一行 JSON（含换行符）"""
    return json.dumps(obj, ensure_ascii=False) + "\n"


def _chat_kwargs_with_thinking(
    disable_thinking: bool = False,
    reasoning_effort: str = "low",
) -> Dict[str, Any]:
    """
    统一构造 chat.completions.create 传递的「思考开关 + 思考强度」额外参数。

    设计意图（给后续维护）：
    - 关思考场景：OCR 图片提取、编排器 JSON 输出、文件文本提取。这些场景只要结构化输出，
      不需要推理链，且思考会污染 <think> 标签，所以用 enable_thinking=false。
    - 开思考场景（默认）：快速模式正文生成、专家模式 reply_directly、深度思考。这些场景
      用户希望在前端"思考栏"看到 reasoning_content，所以传递 reasoning_effort=low 提速。
    - reasoning_effort 取值：low / medium / high，由阿里云 qwen3 系端点广泛支持。
    - 透传方式：统一用 extra_body，兼容 AsyncOpenAI 官方 SDK，不依赖 SDK 原生字段。

    :param disable_thinking: True=关闭思考（OCR/编排器用），False=允许输出思考
    :param reasoning_effort: 思考强度，默认 "low" 提速（用户反馈速度慢）
    :return: 可直接 ** 解包到 chat.completions.create() 的 kwargs dict
    """
    extra_body: Dict[str, Any] = {}
    if disable_thinking:
        logger.debug("[思考开关] 关闭 enable_thinking=False")
        extra_body["enable_thinking"] = False
        return {"extra_body": extra_body}
    # 开思考：只传 reasoning_effort（大多数端点会自动允许思考；部分也需要 enable_thinking）
    if reasoning_effort:
        # 双参数双保险：enable_thinking=true + reasoning_effort，保证各版本端点都能开启且强度生效
        extra_body["enable_thinking"] = True
        extra_body["reasoning_effort"] = reasoning_effort
        logger.debug("[思考开关] 开启 reasoning_effort=%s", reasoning_effort)
    return {"extra_body": extra_body} if extra_body else {}


def _build_thinking_history_prompt(state: ExpertState) -> str:
    """
    构造「思考历史」段落，喂给下轮编排器，让它知道自己之前是怎么想的。
    【硬编码移除】原直接散落字符串拼接；统一走 config.prompts.context.build_thinking_history_text，
    确保多处调用（专家编排器 / deep_thinking 调试）格式完全一致。
    """
    return build_thinking_history_text(state.orch_history)


def _build_global_context(state: ExpertState) -> str:
    """
    组装 FINAL 阶段（deep_thinking / reply_model）用的全局上下文：
    用户提问 + 媒体提取 + 记忆 + 搜索。
    【硬编码移除】
      - 4 个段落标题 / 顺序 / [:3000] 截断，统一由 build_global_context_text 控制，
        保证专家与快速路径（如有）的上下文拼接结果完全一致，避免输出差异。
    """
    return build_global_context_text(
        user_message=state.message,
        extracted_text=state.extracted_text,
        memory_context_parts=state.memory_context_parts,
        search_context_parts=state.all_search_context_parts,
        extracted_max_len=MEDIA_EXTRACTED_MAX_LEN,
    )


class ChatRequest(BaseModel):
    message: str
    history: Optional[List[dict]] = None
    user_id: Optional[int] = None
    session_id: Optional[str] = None
    message_type: Optional[str] = "TEXT"
    media_url: Optional[str] = None
    media_urls: Optional[List[str]] = None
    # 双 UUID（由后端 ChatController 生成）：
    #  - message_uuid         : 用户消息 UUID（role=user）→ 用于更新用户消息正文/提取文本
    #  - assistant_message_uuid: assistant 消息 UUID（role=assistant）→ 用于更新 assistant 侧 searchResults/expertTrace
    # 修复之前共用一个 UUID 导致回写时用户消息和 assistant 消息互相覆盖的 Bug
    message_uuid: Optional[str] = None
    assistant_message_uuid: Optional[str] = None
    mode: Optional[str] = "fast"  # fast: 快速模式, expert: 专家模式


class SummarizeRequest(BaseModel):
    messages: Optional[List[dict]] = None
    existing_summary: Optional[str] = None


async def download_and_extract_docx(media_url: str) -> str:
    # 文件提取是每媒体消息/每批图片的常规动作；批量时「开始+完成」会按文件数打印，统一降到 DEBUG
    logger.debug("[文件提取] 开始下载并提取DOCX文件: %s...", (media_url[:50] if media_url else ""))

    try:
        import tempfile
        import os

        async with httpx.AsyncClient() as http_client:
            response = await http_client.get(media_url)
            response.raise_for_status()

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(response.content)
                temp_file_path = temp_file.name

        try:
            from docx import Document
            doc = Document(temp_file_path)

            all_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    all_content.append(paragraph.text)

            for table in doc.tables:
                table_text = ["表格:"]
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if len(table_text) > 1:
                    all_content.append("\n".join(table_text))

            extracted_text = "\n".join(all_content)
            logger.debug("[文件提取] DOCX提取完成，段落数=%s, 表格数=%s, 总字符数=%s",
                         len(doc.paragraphs), len(doc.tables), len(extracted_text))

            # 【硬编码移除】[:3000] → MEDIA_EXTRACTED_MAX_LEN（与 PDF/TXT/OCR 多图共用同一长度）
            return extracted_text[:MEDIA_EXTRACTED_MAX_LEN]
        except ImportError:
            logger.warning("[文件提取] python-docx库未安装，无法提取DOCX内容")
            return ""
        except Exception as e:
            logger.error("[文件提取] 提取DOCX内容失败: %s", str(e))
            return ""
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
    except Exception as e:
        logger.error("[文件提取] 下载DOCX文件失败: %s", str(e))
        return ""


async def download_and_extract_pdf(media_url: str) -> str:
    logger.debug("[文件提取] 开始下载并提取PDF文件: %s...", (media_url[:50] if media_url else ""))

    try:
        import tempfile
        import os

        async with httpx.AsyncClient() as http_client:
            response = await http_client.get(media_url)
            response.raise_for_status()

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_file:
                temp_file.write(response.content)
                temp_file_path = temp_file.name

        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(temp_file_path)

            all_content = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    all_content.append(page_text)

            extracted_text = "\n\n".join(all_content)
            logger.debug("[文件提取] PDF提取完成，页数=%s, 总字符数=%s", len(reader.pages), len(extracted_text))

            # 【硬编码移除】[:3000] → MEDIA_EXTRACTED_MAX_LEN
            return extracted_text[:MEDIA_EXTRACTED_MAX_LEN]
        except ImportError:
            logger.warning("[文件提取] PyPDF2库未安装，无法提取PDF内容")
            return ""
        except Exception as e:
            logger.error("[文件提取] 提取PDF内容失败: %s", str(e))
            return ""
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
    except Exception as e:
        logger.error("[文件提取] 下载PDF文件失败: %s", str(e))
        return ""


async def download_and_read_text(media_url: str) -> str:
    logger.debug("[文件提取] 开始下载并读取文本文件: %s...", (media_url[:50] if media_url else ""))

    try:
        async with httpx.AsyncClient() as http_client:
            response = await http_client.get(media_url)
            response.raise_for_status()

            text = response.text
            logger.debug("[文件提取] 文本文件读取完成，总字符数=%s", len(text))

            # 【硬编码移除】[:3000] → MEDIA_EXTRACTED_MAX_LEN
            return text[:MEDIA_EXTRACTED_MAX_LEN]
    except Exception as e:
        logger.error(f"[文件提取] 下载或读取文本文件失败: {str(e)}")
        return ""


async def extract_media_text(media_url: str, media_type: str) -> str:
    start_time = time.time()

    # 【硬编码移除】"IMAGE" / "FILE" 字面量 → MEDIA_TYPE_IMAGE / MEDIA_TYPE_FILE 协议常量
    if media_type == MEDIA_TYPE_IMAGE:
        if ocr_client:
            client = ocr_client
            model = OCR_MODEL
            # 选择哪个模型来提取：DEBUG 即可（每媒体消息都会打一次，重复度高）
            logger.debug("[图片提取] 使用OCR模型: %s", OCR_MODEL)
        elif multimodal_client:
            client = multimodal_client
            model = MULTIMODAL_MODEL
            logger.debug("[图片提取] 使用多模态模型: %s", MULTIMODAL_MODEL)
        else:
            logger.error("[图片提取] OCR和多模态客户端均未初始化，无法提取图片内容")
            return ""

        # 【硬编码移除】单图 OCR 的 system_prompt / 用户侧文本 统一读 config.prompts.ocr
        system_prompt = OCR_SINGLE_IMAGE_SYSTEM_PROMPT

        # 【硬编码移除】URL 前缀字面量 "data:image" / "http" → 常量 URL_PREFIX_DATA_IMAGE / URL_PREFIX_HTTP
        if media_url.startswith(URL_PREFIX_DATA_IMAGE):
            content = [
                {"type": "text", "text": OCR_SINGLE_IMAGE_USER_TEXT},
                {"type": "image_url", "image_url": {"url": media_url}},
            ]
        elif media_url.startswith(URL_PREFIX_HTTP):
            content = [
                {"type": "text", "text": OCR_SINGLE_IMAGE_USER_TEXT},
                {"type": "image_url", "image_url": {"url": media_url}},
            ]
        else:
            logger.error("[图片提取] 不支持的图片URL格式: %s...", (media_url[:30] if media_url else ""))
            return ""

        try:
            # OCR 提取必须关闭思考：否则模型先输出 reasoning 会污染提取到的纯文本
            # 【硬编码移除】temperature=0.3 → OCR_TEMPERATURE；关思考复用 OCR_DISABLE_THINKING
            response = await client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": content}
                ],
                stream=False,
                temperature=OCR_TEMPERATURE,
                **_chat_kwargs_with_thinking(disable_thinking=bool(OCR_DISABLE_THINKING))
            )

            duration = time.time() - start_time

            if response.usage:
                logger.debug("[图片提取] Token消耗 model=%s prompt=%s completion=%s total=%s dur=%.2fs",
                             model, response.usage.prompt_tokens, response.usage.completion_tokens,
                             response.usage.total_tokens, duration)

            if response.choices and response.choices[0].message and response.choices[0].message.content:
                extracted_text = response.choices[0].message.content.strip()
                # 单图提取成功：DEBUG，[媒体提取] 汇总行已经打 INFO 了
                logger.debug("[图片提取] 成功 model=%s len=%s dur=%.2fs", model, len(extracted_text), duration)
                return extracted_text
            else:
                logger.warning("[图片提取] 返回空内容 dur=%.2fs", duration)
                return ""
        except Exception as e:
            duration = time.time() - start_time
            logger.error("[图片提取] 失败 model=%s err=%s dur=%.2fs", model, str(e), duration)
            return ""
    else:  # FILE
        if not media_url.startswith(URL_PREFIX_HTTP):
            logger.error("[文件提取] 不支持的文件URL格式: %s...", (media_url[:30] if media_url else ""))
            return ""

        url_lower = media_url.lower()

        # 【硬编码移除】原 if/elif 散落 ".docx"/".pdf"/".txt" 等；改用配置白名单 SUPPORTED_FILE_EXTENSIONS / PLAIN_TEXT_EXTENSIONS
        if not url_lower.endswith(SUPPORTED_FILE_EXTENSIONS):
            logger.warning("[文件提取] 不支持的文件扩展名: %s", url_lower.split("?")[0][-30:])
            return ""

        if url_lower.endswith(".docx"):
            return await download_and_extract_docx(media_url)
        elif url_lower.endswith(".pdf"):
            return await download_and_extract_pdf(media_url)
        # 【硬编码移除】".txt"/".md"/".json"/".csv"/".xml" 分散字面量 → PLAIN_TEXT_EXTENSIONS（behavior.py 单源）
        elif url_lower.endswith(PLAIN_TEXT_EXTENSIONS):
            return await download_and_read_text(media_url)
        else:
            logger.warning("[文件提取] 不支持的文件类型: %s...", (media_url[:50] if media_url else ""))
            return ""


async def extract_media_text_batch(media_urls: List[str], media_type: str) -> str:
    """批量提取媒体文本 - 使用并行处理提高速度。
    【硬编码移除】
      - "IMAGE" 字面量 → MEDIA_TYPE_IMAGE
    """
    start_time = time.time()
    # 批量入口降 DEBUG：外层 [请求开始] 已记录 type/count，批量开始 INFO 价值低
    logger.debug("[批量提取] 开始 type=%s count=%s", media_type, len(media_urls))

    if media_type == MEDIA_TYPE_IMAGE and len(media_urls) > 1:
        # 图片类型：尝试一次调用处理多张图片（支持多图输入的模型）
        extracted_text = await extract_multiple_images(media_urls)
        if extracted_text:
            duration = time.time() - start_time
            # 批量路径成功：DEBUG
            logger.debug("[批量提取] 批量处理成功 len=%s dur=%.2fs", len(extracted_text), duration)
            return extracted_text
        else:
            # 批量处理失败，回退到并行处理
            logger.warning("[批量提取] 批量处理失败，回退到并行处理")

    # 并行处理多个媒体文件
    tasks = []
    for i, url in enumerate(media_urls):
        # 启动子任务日志降到 DEBUG
        logger.debug("[批量提取] 启动任务 %s/%s: %s...", i + 1, len(media_urls), (url[:50] if url else ""))
        tasks.append(extract_media_text(url, media_type))

    # 使用asyncio.gather并发执行
    results = await asyncio.gather(*tasks)

    all_extracted = []
    for i, text in enumerate(results):
        if text:
            all_extracted.append(f"图{i+1}：{text}")
        else:
            all_extracted.append(f"图{i+1}：[提取失败]")

    duration = time.time() - start_time
    result = "\n\n".join(all_extracted)
    # 并行完成降 DEBUG：外层 [请求结束] 汇总行能体现总耗时
    logger.debug("[批量提取] 并行完成 type=%s total_len=%s dur=%.2fs", media_type, len(result), duration)

    return result


async def extract_multiple_images(media_urls: List[str]) -> str:
    """一次调用处理多张图片 - 支持多图输入的模型。
    【硬编码移除】
      - system_prompt / 用户侧文本 → config.prompts.ocr 的多图版本
      - temperature=0.3 → OCR_TEMPERATURE；关思考 → OCR_DISABLE_THINKING
    """
    start_time = time.time()
    # 多图提取启动日志降 DEBUG；批量提取入口已打 INFO 一行
    logger.debug("[多图提取] 一次调用处理 %s 张图片", len(media_urls))

    if ocr_client:
        client = ocr_client
        model = OCR_MODEL
        logger.debug("[多图提取] 使用OCR模型: %s", OCR_MODEL)
    elif multimodal_client:
        client = multimodal_client
        model = MULTIMODAL_MODEL
        logger.debug("[多图提取] 使用多模态模型: %s", MULTIMODAL_MODEL)
    else:
        logger.error("[多图提取] OCR和多模态客户端均未初始化")
        return ""

    # 【硬编码移除】多图 OCR 的大段 system / user 文本 → 配置文件
    system_prompt = OCR_MULTI_IMAGE_SYSTEM_PROMPT

    # 构建多图内容
    content = [{"type": "text", "text": OCR_MULTI_IMAGE_USER_TEXT}]
    for i, url in enumerate(media_urls):
        content.append({"type": "image_url", "image_url": {"url": url, "index": i + 1}})

    try:
        # 多图提取必须关闭思考：reasoning 内容会混入每张图片的分析正文，破坏"图N："格式
        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": content}
            ],
            stream=False,
            temperature=OCR_TEMPERATURE,
            **_chat_kwargs_with_thinking(disable_thinking=bool(OCR_DISABLE_THINKING))
        )

        duration = time.time() - start_time

        if response.usage:
            logger.debug("[多图提取] Token消耗 model=%s prompt=%s completion=%s total=%s dur=%.2fs",
                         model, response.usage.prompt_tokens, response.usage.completion_tokens,
                         response.usage.total_tokens, duration)

        if response.choices and response.choices[0].message and response.choices[0].message.content:
            extracted_text = response.choices[0].message.content.strip()
            logger.debug("[多图提取] 成功 model=%s len=%s dur=%.2fs", model, len(extracted_text), duration)
            return extracted_text
        else:
            logger.warning("[多图提取] 返回空内容 dur=%.2fs", duration)
            return ""
    except Exception as e:
        duration = time.time() - start_time
        logger.error("[多图提取] 失败 model=%s err=%s dur=%.2fs", model, str(e), duration)
        return ""


async def stream_model_response(message: str, history: Optional[List[dict]] = None,
                                user_id: Optional[int] = None, session_id: Optional[int] = None,
                                message_type: str = "TEXT", media_url: Optional[str] = None,
                                media_urls: Optional[List[str]] = None,
                                message_uuid: Optional[str] = None,
                                mode: str = "fast",
                                assistant_message_uuid: Optional[str] = None):
    request_start_time = time.time()
    # 紧凑的一行汇总 INFO（不再打 4~6 行 + 分隔线），避免每次请求头尾 10+ 行分隔线刷屏
    # 打印两个 UUID 便于和后端回写日志对齐
    logger.info(
        "[请求开始] mode=%s type=%s session=%s user=%s history=%s userUuid=%s assistantUuid=%s msg=%s",
        mode, message_type, session_id, user_id,
        (len(history) if history else 0),
        (message_uuid or "")[:8] + "…" if message_uuid else "-",
        (assistant_message_uuid or "")[:8] + "…" if assistant_message_uuid else "-",
        # 【硬编码移除】[:100] → LOG_USER_MESSAGE_TRUNCATE_LEN（请求入口日志统一截断长度）
        (message[:LOG_USER_MESSAGE_TRUNCATE_LEN] +
         ("…" if message and len(message) > LOG_USER_MESSAGE_TRUNCATE_LEN else "")),
    )
    
    # 根据模式选择处理函数
    if mode == "expert":
        async for result in expert_mode_process(message, history, user_id, session_id, 
                                                message_type, media_url, media_urls, 
                                                message_uuid, assistant_message_uuid,
                                                request_start_time):
            yield result
        return

    extracted_text = ""
    # 【硬编码移除】"IMAGE"/"FILE" → 协议常量
    is_multimodal = message_type in (MEDIA_TYPE_IMAGE, MEDIA_TYPE_FILE)
    
    if is_multimodal and (media_url or (media_urls and len(media_urls) > 0)):
        if media_urls and len(media_urls) > 0:
            extracted_text = await extract_media_text_batch(media_urls, message_type)
        else:
            extracted_text = await extract_media_text(media_url, message_type)
        
        if extracted_text:
            # 媒体提取成功降 DEBUG：[请求结束] 汇总有 reply_len，结合请求的 message_type 能判断是否提取成功
            logger.debug("[媒体提取] 成功，长度: %s字符", len(extracted_text))
        else:
            logger.warning("[媒体提取] 失败，将使用原始消息进行记忆检索")

    memories = []
    if session_id:
        mem_start = time.time()
        try:
            search_query = (
                f"{message} {extracted_text}" if extracted_text
                else (message if message_type in ("TEXT", "VOICE")
                      else (f"{message} {media_url}" if media_url else message))
            )
            # 【硬编码移除】top_k=5 → MEMORY_SEARCH_DEFAULT_TOP_K（与专家工具保持一致）
            memories = await memory_service.search_memories(
                search_query, session_id, top_k=MEMORY_SEARCH_DEFAULT_TOP_K
            )
            # 记忆检索：DEBUG 即可（失败仍保留 ERROR）
            logger.debug("[记忆检索] 完成 count=%s dur=%.2fs", len(memories), time.time() - mem_start)
        except Exception as e:
            logger.error("[记忆检索] 失败: %s", e)

    search_results = []
    search_context = ""
    # 累积工具调用记录，用于持久化到 assistant 消息的 expertTrace 字段
    fast_mode_tool_history: List[Dict[str, Any]] = []

    context_text = message
    if extracted_text:
        context_text += f"\n{extracted_text}"
    if memories:
        memories_text = "\n".join([f"- {m['content']}" for m in memories])
        context_text += f"\n\n相关记忆：\n{memories_text}"

    orch_start = time.time()
    # 流式调用编排器：边生成边 yield analysis 增量给前端，实现"打字机"效果
    orch_result_holder = {}
    async for analysis_delta in _call_orchestration_model_stream(context_text, message, orch_result_holder):
        # 前端收到 orchestration_chunk 后实时显示编排分析文本（逐步增长）
        yield _json_line({"type": "orchestration_chunk", "data": {"delta": analysis_delta}})
    # 编排器耗时：DEBUG 即可；结果在下面一行打 INFO
    logger.debug("[编排器] 调用完成 dur=%.2fs", time.time() - orch_start)

    orchestration_result = orch_result_holder.get("result") or {"need_search": False, "search_keywords": [], "analysis_text": "编排器调用失败"}

    if orchestration_result:
        need_search = orchestration_result.get("need_search", False)
        search_keywords = orchestration_result.get("search_keywords", [])
        analysis_text = orchestration_result.get("analysis_text", "")

        # 编排器结果降 DEBUG：[请求结束] 已有 search=X 条数，失败仍 ERROR；需要看决策细节开 DEBUG
        logger.debug("[编排器结果] need_search=%s keywords=%s", need_search, search_keywords)

        # ===== 编排步骤完成：通知前端此步的结构化信息（analysis 已通过 chunk 流式显示）=====
        yield _json_line({"type": "orchestration_step", "data": {
            "iteration": 1,
            # 【硬编码移除】"planning" → ORCH_PHASE_PLANNING
            "phase": ORCH_PHASE_PLANNING,
            "action": "search" if need_search else "reply",
            "purpose": None,
            # analysis 不再重复发送（已通过 orchestration_chunk 逐步显示给前端）
        }})

        if need_search and search_keywords:
            # ===== 新增：工具调用开始通知（替代旧 search_start）=====
            yield _json_line({"type": "tool_call_start", "data": {
                "tool": TOOL_NAME_WEB_SEARCH,
                "params": {"keywords": search_keywords},
            }})
            search_start = time.time()
            search_results = await search_service.web_search(search_keywords)
            search_context = await search_service.get_search_context(search_keywords)
            search_dur_ms = int((time.time() - search_start) * 1000)
            # 过滤结果（只保留 title/url）
            filtered_results = [
                {"title": r.get("title", ""), "url": r.get("url", "")}
                for r in (search_results or [])
                if r.get("url")
            ]
            # ===== 新增：工具调用完成通知（替代旧 search_summary，前端据此渲染工具摘要）=====
            yield _json_line({"type": "tool_call_result", "data": {
                "tool": TOOL_NAME_WEB_SEARCH,
                "success": True,
                "durationMs": search_dur_ms,
                "resultCount": len(filtered_results),
                "summary": SUMMARY_TEMPLATE_WEB_SEARCH.format(n=len(filtered_results)),
                "results": filtered_results,
            }})
            # 快速模式联网搜索：DEBUG 即可；结果条数在 [请求结束] 汇总行已经有了
            logger.debug("[联网搜索] 完成 count=%s dur=%.2fs", len(search_results), search_dur_ms / 1000.0)
            # 累积工具调用记录，用于持久化到 assistant 消息
            fast_mode_tool_history.append({
                "iteration": 1,
                # 【硬编码移除】动作名 "collect_tools" → EXPERT_ACTION_COLLECT（与状态机一致）
                "action": EXPERT_ACTION_COLLECT,
                "purpose": "联网搜索",
                "analysis": orchestration_result.get("analysis_text", ""),
                "tools": [
                    {
                        "tool": TOOL_NAME_WEB_SEARCH,
                        "params": {"keywords": search_keywords},
                        "success": True,
                        "durationMs": search_dur_ms,
                        "resultCount": len(filtered_results),
                        "error": None,
                        "rawResult": filtered_results,
                    }
                ],
            })

    # 【硬编码移除】主模型人设 system prompt 统一读 persona.py
    system_prompt = PERSONA_SYSTEM_PROMPT

    if memories:
        memories_text = "\n".join([f"- {m['content']}" for m in memories])
        # 【硬编码移除】"\n\n用户长期记忆：\n" 段落标题 → PERSONA_MEMORY_SECTION_TITLE
        system_prompt += PERSONA_MEMORY_SECTION_TITLE + memories_text

    if search_context:
        # 【硬编码移除】"\n\n联网搜索信息：\n" 段落标题 → PERSONA_SEARCH_SECTION_TITLE
        system_prompt += PERSONA_SEARCH_SECTION_TITLE + search_context

    messages = [{"role": "system", "content": system_prompt}]

    if history:
        for msg in history:
            role = msg.get("role")
            content = msg.get("content")
            msg_type = msg.get("messageType", "TEXT")
            extracted_text_history = msg.get("extractedText")
            
            if role and content:
                # 【硬编码移除】历史消息前缀字符串 → config.prompts.context.*_PREFIX 常量
                if msg_type == MEDIA_TYPE_IMAGE:
                    if extracted_text_history:
                        content = HISTORY_IMAGE_PREFIX_WITH_EXTRACTED.format(
                            content=content, extracted_text=extracted_text_history
                        )
                    else:
                        content = HISTORY_IMAGE_PREFIX_SIMPLE.format(content=content)
                elif msg_type == MEDIA_TYPE_FILE:
                    if extracted_text_history:
                        content = HISTORY_FILE_PREFIX_WITH_EXTRACTED.format(
                            content=content, extracted_text=extracted_text_history
                        )
                    else:
                        content = HISTORY_FILE_PREFIX_SIMPLE.format(content=content)
                messages.append({"role": role, "content": content})

    # 【硬编码移除】mediaType 判断 → MEDIA_TYPE_IMAGE / MEDIA_TYPE_FILE
    is_multimodal = message_type == MEDIA_TYPE_IMAGE
    
    if is_multimodal and (media_urls and len(media_urls) > 0):
        user_content = [{"type": "text", "text": message}]
        for url in media_urls:
            user_content.append({"type": "image_url", "image_url": {"url": url}})
    elif is_multimodal and media_url:
        user_content = [
            {"type": "text", "text": message},
            {"type": "image_url", "image_url": {"url": media_url}}
        ]
    elif message_type == MEDIA_TYPE_FILE and (media_url or (media_urls and len(media_urls) > 0)):
        # 【硬编码移除】文件内容前缀 → FILE_TEXT_BLOCK_PREFIX；单/多文件整行模板 → USER_PREFIX_UPLOAD_SINGLE_FILE / _MULTI_FILES
        file_text = (FILE_TEXT_BLOCK_PREFIX.format(text=extracted_text) if extracted_text else "")
        if media_urls and len(media_urls) > 0:
            user_content = USER_PREFIX_UPLOAD_MULTI_FILES.format(
                n_files=len(media_urls), message=message, file_text_block=file_text
            )
        else:
            user_content = USER_PREFIX_UPLOAD_SINGLE_FILE.format(
                media_url=media_url, message=message, file_text_block=file_text
            )
    else:
        user_content = message

    # 单图快速模式（只有 1 张 IMAGE URL）→ 用模板前缀包装 user_content 便于人设上下文与 OCR 一致
    if (message_type == MEDIA_TYPE_IMAGE
            and media_url
            and not (media_urls and len(media_urls) > 0)
            and isinstance(user_content, str)):
        if extracted_text:
            user_content = USER_PREFIX_UPLOAD_SINGLE_IMAGE_WITH_EXTRACTED.format(
                message=message, extracted_text=extracted_text
            )
        else:
            user_content = USER_PREFIX_UPLOAD_SINGLE_IMAGE.format(message=message)
    
    messages.append({"role": "user", "content": user_content})
    # 消息构建：DEBUG 即可（因为 [请求开始]/[请求结束] 两行汇总已经能定位问题）
    logger.debug("[消息构建] messages=%s user_content_len=%s", len(messages), len(str(user_content)))

    assistant_response = ""
    if is_multimodal:
        current_model = MULTIMODAL_MODEL
        current_client = multimodal_client
    elif message_type == "FILE":
        current_model = MODEL
        current_client = client
    else:
        current_model = MODEL
        current_client = client

    try:
        model_start = time.time()

        # 快速模式正文：
        # 用户明确要求「只有深度思考模型开思考」，快速模式（聊天主流程）统一关闭思考。
        # 之前临时开过 reasoning_effort=low 用于测试思考栏，现在回滚。
        # 关思考的角色：OCR/编排器/文件提取 + 快速模式正文 + 专家模式简单回复；
        # 开思考的角色：仅深度思考模型（_expert_final_deep_thinking 内显式开启）。
        disable_thinking_for_fast = True
        logger.debug("[快速模式正文] 模型=%s, 关思考=%s", current_model, disable_thinking_for_fast)

        response = await current_client.chat.completions.create(
            model=current_model,
            messages=messages,
            stream=True,
            **_chat_kwargs_with_thinking(
                disable_thinking=disable_thinking_for_fast,
                reasoning_effort="low",
            )
        )
        # 成功连接：DEBUG（[流式传输] 完成的汇总行已经能证明 API 连通且开始/结束）
        logger.debug("[模型调用] 成功连接到模型API model=%s", current_model)

        chunk_count = 0
        empty_delta_count = 0
        total_content = ""
        # 新增：累计 reasoning_content（阿里云 qwen3.8 系模型 chunk 会先推 reasoning_delta），
        # 拿到 reasoning 就立即 yield 一条 type=thinking 给前端渲染到黄色思考栏，不等正文。
        total_reasoning = ""
        async for chunk in response:
            if chunk.choices and chunk.choices[0].delta:
                delta = chunk.choices[0].delta
                # 先处理 reasoning：字段名兼容多种（reasoning_content 是官方命名，
                # thinking_content / reasoning_delta 是不同版本端点的别名）
                reasoning = (
                    getattr(delta, "reasoning_content", None)
                    or getattr(delta, "thinking_content", None)
                    or getattr(delta, "reasoning_delta", None)
                )
                if reasoning:
                    total_reasoning += reasoning
                    yield _json_line({"type": "thinking", "data": reasoning})
                content = delta.content
                if content:
                    chunk_count += 1
                    total_content += content
                    assistant_response += content
                    # 每个 chunk 的原文只打 DEBUG（之前默认 INFO 级会每收到一字就刷屏）
                    logger.debug("[流式传输] 分片#%s content_len=%s", chunk_count, len(content))
                    yield _json_line({"type": "content", "data": content})
                else:
                    # 空 delta 也可能带 finish_reason，只计数，不逐条打印
                    empty_delta_count += 1
            else:
                empty_delta_count += 1

        model_duration = time.time() - model_start
        
        if hasattr(response, 'usage') and response.usage:
            # Token消耗：保留 INFO（计费/性能排查核心证据）
            logger.info(
                "[模型调用] Token消耗 model=%s prompt=%s completion=%s total=%s dur=%.2fs",
                current_model,
                response.usage.prompt_tokens, response.usage.completion_tokens,
                response.usage.total_tokens, model_duration,
            )
        
        # 流式结束聚合：降 DEBUG（[请求结束] 汇总行已包含 reply_len + total dur）
        # 补充 reasoning_len：方便判断本次模型到底有没有输出 reasoning
        logger.debug(
            "[流式传输] 完成 chunks=%s empty=%s content_len=%s reasoning_len=%s model=%s dur=%.2fs",
            chunk_count, empty_delta_count, len(total_content), len(total_reasoning),
            current_model, model_duration,
        )

        if chunk_count == 0:
            logger.warning("[流式传输] 未收到模型返回的内容")

        # 注意：搜索结果已在工具调用时通过 tool_call_result 块实时返回，不再在末尾重复汇总
        # search_results 变量保留给 [请求结束] 汇总行使用

    except Exception as e:
        logger.error(f"[模型调用] 出错: {str(e)}")
        yield f"错误: {str(e)}"

    # 回写图片提取内容到用户消息（必须用 user_message_uuid，不能误用 assistant UUID）
    if user_id and assistant_response:
        # 修复：多图时 media_url 为 None 但 media_urls 有值，条件应同时检查两者
        if extracted_text and (media_url or (media_urls and len(media_urls) > 0)):
            combined_content = f"[用户提问]{message}\n[图片内容]{extracted_text}" if message_type == "IMAGE" else f"[用户提问]{message}\n[文件内容]{extracted_text}"
            
            asyncio.create_task(
                memory_service.async_extract_and_store(
                    combined_content, assistant_response, user_id, session_id
                )
            )
            
            if message_uuid:
                await update_backend_message_content(session_id, message_uuid, message, extracted_text)
        else:
            asyncio.create_task(
                memory_service.async_extract_and_store(
                    message, assistant_response, user_id, session_id
                )
            )

    # 通过 SSE 流发送工具调用记录（替代 HTTP 回调，避免竞态）
    # 后端在 SSE 流中捕获 expert_trace 事件，流结束时随 saveMessage 一起持久化
    if fast_mode_tool_history:
        fast_trace_payload = {
            # 【硬编码移除】"fast_reply" → FAST_MODE_FINAL_PATH
            "finalPath": FAST_MODE_FINAL_PATH,
            "iterationCount": len(fast_mode_tool_history),
            "history": fast_mode_tool_history,
            "deepThinkingReasoning": "",
        }
        yield _json_line({
            "type": "expert_trace",
            "data": json.dumps(fast_trace_payload, ensure_ascii=False),
        })

    total_duration = time.time() - request_start_time
    # 请求结束：一行汇总，替代原来的 6 行分隔线 + 多行分类
    logger.info(
        "[请求结束] dur=%.2fs mode=%s reply_len=%s search=%s mem=%s session=%s",
        total_duration, mode, len(assistant_response),
        len(search_results), len(memories), session_id,
    )


async def _call_orchestration_model_stream(
    context_text: str, message: str, result_holder: dict
) -> AsyncGenerator[str, None]:
    """
    流式调用快速模式编排器：判断 need_search + 输出 search_keywords。
    【硬编码移除】
      - system_prompt / user_prompt 统一读 config.prompts.orchestrator_fast
      - 上下文截断 [:2000] → FAST_ORCH_CONTEXT_MAX_LEN
      - max_retries / temperature → config.behavior FAST_ORCH_*
    """
    if not orchestration_client:
        logger.warning("[编排器] client未初始化，跳过编排")
        result_holder["error"] = "编排器未初始化"
        return
    
    # 【硬编码移除】system_prompt → FAST_ORCH_SYSTEM_PROMPT
    system_prompt = FAST_ORCH_SYSTEM_PROMPT

    # 【硬编码移除】首次 user_prompt 走 build_fast_orch_user_prompt（上下文截断 FAST_ORCH_CONTEXT_MAX_LEN）
    user_prompt = build_fast_orch_user_prompt(
        message, context_text[:FAST_ORCH_CONTEXT_MAX_LEN], enforce_strict=False
    )

    # 【硬编码移除】max_retries=2 → FAST_ORCH_MAX_RETRIES
    max_retries = FAST_ORCH_MAX_RETRIES
    # 第一次尝试用流式，后续重试改非流式（避免前端重复显示 analysis 文本）
    use_stream = True
    full_text = ""
    last_analysis_len = 0

    for attempt in range(1, max_retries + 1):
        try:
            response = await orchestration_client.chat.completions.create(
                model=ORCHESTRATION_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                stream=use_stream,
                # 【硬编码移除】temperature=0.3 → FAST_ORCH_TEMPERATURE
                temperature=FAST_ORCH_TEMPERATURE,
                # 编排器输出 JSON，必须关思考：否则 reasoning 文本会混入 JSON 导致解析失败
                **_chat_kwargs_with_thinking(disable_thinking=True),
            )

            if use_stream:
                # ===== 流式模式：边接收边提取 analysis_text 增量 =====
                full_text = ""
                last_analysis_len = 0
                async for chunk in response:
                    if not chunk.choices or not chunk.choices[0].delta:
                        continue
                    delta = chunk.choices[0].delta.content
                    if not delta:
                        continue
                    full_text += delta
                    # 增量提取 analysis_text：只返回比上次多出来的部分
                    current = _try_extract_analysis(full_text, keys=("analysis_text",))
                    if len(current) > last_analysis_len:
                        yield current[last_analysis_len:]
                        last_analysis_len = len(current)
            else:
                # ===== 非流式重试模式 =====
                if response.usage:
                    logger.info(f"[编排器] Token消耗(模型:{ORCHESTRATION_MODEL}): prompt={response.usage.prompt_tokens}, completion={response.usage.completion_tokens}, total={response.usage.total_tokens}")
                if response.choices and response.choices[0].message and response.choices[0].message.content:
                    full_text = response.choices[0].message.content.strip()
                else:
                    full_text = ""

            # 解析完整 JSON
            cleaned = _clean_json_response(full_text)
            try:
                result = json.loads(cleaned)
                if isinstance(result, dict):
                    if not use_stream:
                        logger.info(f"[编排器] 重试成功")
                    result_holder["result"] = {
                        "need_search": result.get("need_search", False),
                        "search_keywords": result.get("search_keywords", []),
                        "analysis_text": result.get("analysis_text", ""),
                    }
                    return
            except json.JSONDecodeError as je:
                logger.warning(f"[编排器] 第{attempt}次调用JSON解析失败: {je}")
                if attempt < max_retries:
                    # 【硬编码移除】重试 user_prompt → enforce_strict=True 模式（会追加强约束行）；上下文截断仍走 FAST_ORCH_CONTEXT_MAX_LEN
                    user_prompt = build_fast_orch_user_prompt(
                        message, context_text[:FAST_ORCH_CONTEXT_MAX_LEN], enforce_strict=True
                    )
                    use_stream = False
                    continue
        except Exception as e:
            logger.error(f"[编排器] 第{attempt}次调用失败({ORCHESTRATION_MODEL}): {e}")
            if attempt < max_retries:
                use_stream = False
                continue

    logger.error(f"[编排器] 调用失败({ORCHESTRATION_MODEL})，已重试{max_retries}次")
    result_holder["error"] = "编排器调用失败"


def _try_extract_analysis(full_text: str, keys: tuple = ("analysis_text", "analysis")) -> str:
    """
    从可能不完整的流式 JSON 文本中提取 analysis/analysis_text 字段的值。

    使用正则匹配 JSON 字符串值（支持转义），然后用 json.loads 反转义。
    如果字段值尚未闭合（流式传输中），返回到当前末尾的所有内容。

    Args:
        full_text: 到目前为止累积的完整模型输出文本
        keys: 要查找的字段名元组，按优先级排序

    Returns: 已反转义的 analysis 文本（可能不完整）；如果未找到字段则返回空字符串
    """
    import re

    for key in keys:
        # 构造正则：匹配 "key" : "value" 中的 value 部分（支持转义字符）
        # (?:[^"\\]|\\.)* 匹配：非引号非反斜杠的任意字符，或反斜杠+任意字符（转义序列）
        key_quoted = re.escape(f'"{key}"')
        # 先尝试匹配闭合的字符串值
        pattern_closed = key_quoted + r'\s*:\s*"((?:[^"\\]|\\.)*)"'
        m = re.search(pattern_closed, full_text, re.DOTALL)
        if m:
            raw = m.group(1)
            try:
                return json.loads('"' + raw + '"')
            except (json.JSONDecodeError, ValueError):
                # json.loads 失败时做简单反转义兜底
                return raw.replace('\\n', '\n').replace('\\"', '"').replace('\\\\', '\\')

        # 再尝试匹配未闭合的字符串值（流式传输中间状态）
        pattern_open = key_quoted + r'\s*:\s*"((?:[^"\\]|\\.)*)'
        m = re.search(pattern_open, full_text, re.DOTALL)
        if m:
            raw = m.group(1)
            try:
                return json.loads('"' + raw + '"')
            except (json.JSONDecodeError, ValueError):
                return raw.replace('\\n', '\n').replace('\\"', '"').replace('\\\\', '\\')

    return ""


def _clean_json_response(content: str) -> str:
    if not content:
        return content
    
    content = content.strip()
    
    if content.startswith("```"):
        first_backtick = content.find("```")
        last_backtick = content.rfind("```")
        if first_backtick != last_backtick:
            content = content[first_backtick + 3:last_backtick]
        else:
            content = content[3:]
    
    content = content.strip()
    if content.startswith("json"):
        content = content[4:].strip()
    
    import re
    json_match = re.search(r'\{[\s\S]*\}', content)
    if json_match:
        content = json_match.group(0)
    
    return content.strip()


async def update_backend_search_results(session_id: int, message_uuid: str, search_results: list):
    # 【硬编码移除】回写接口路径统一用 BACKEND_ENDPOINT_UPDATE_SEARCH_RESULTS；
    # timeout=30.0 → BACKEND_INTERNAL_API_TIMEOUT + BACKEND_INTERNAL_API_CONNECT_TIMEOUT
    update_url = f"{BACKEND_BASE_URL}{BACKEND_ENDPOINT_UPDATE_SEARCH_RESULTS}"

    try:
        search_results_json = json.dumps(search_results, ensure_ascii=False)

        async with httpx.AsyncClient(timeout=httpx.Timeout(
            BACKEND_INTERNAL_API_TIMEOUT,
            connect=BACKEND_INTERNAL_API_CONNECT_TIMEOUT,
        )) as http_client:
            response = await http_client.post(
                update_url,
                json={
                    "session_id": session_id,
                    "message_uuid": message_uuid,
                    "search_results": search_results_json
                },
                headers={
                    "X-Internal-Secret": INTERNAL_SECRET
                },
            )

            if response.status_code == 200:
                # 内部回调成功：DEBUG（每次搜索成功就一次，和 expertTrace 回写一起，3+条/请求）
                # 失败保留 ERROR
                logger.debug("[搜索结果保存] 成功: session_id=%s, message_uuid=%s, count=%s",
                    session_id, message_uuid, len(search_results))
            else:
                logger.error("[搜索结果保存] 失败: status=%s, %s", response.status_code, response.text)
    except Exception as e:
        logger.error("[搜索结果保存] 调用后端失败: %s", str(e))


async def update_backend_message_content(session_id: int, message_uuid: str, message: str, extracted_text: str):
    # 【硬编码移除】回写接口路径 BACKEND_ENDPOINT_UPDATE_MESSAGE_CONTENT；
    # timeout=30.0 → BACKEND_INTERNAL_API_TIMEOUT + BACKEND_INTERNAL_API_CONNECT_TIMEOUT
    update_url = f"{BACKEND_BASE_URL}{BACKEND_ENDPOINT_UPDATE_MESSAGE_CONTENT}"

    try:
        async with httpx.AsyncClient(timeout=httpx.Timeout(
            BACKEND_INTERNAL_API_TIMEOUT,
            connect=BACKEND_INTERNAL_API_CONNECT_TIMEOUT,
        )) as client:
            response = await client.post(
                update_url,
                json={
                    "session_id": session_id,
                    "message_uuid": message_uuid,
                    "message": message,
                    "extracted_text": extracted_text
                },
                headers={
                    "X-Internal-Secret": INTERNAL_SECRET
                },
            )

            if response.status_code == 200:
                logger.info(f"[消息回写] 成功: session_id={session_id}, message_uuid={message_uuid}")
            else:
                logger.error(f"[消息回写] 失败: status={response.status_code}, {response.text}")
    except Exception as e:
        logger.error(f"[消息回写] 调用后端失败: {str(e)}")


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/chat")
async def chat(request: ChatRequest):
    # 每条请求绑定 correlation_id（自动注入到所有 logger 输出前缀），并发请求日志不串
    req_id = bind_req_id()
    logger.info(
        "[API入口] mode=%s type=%s history=%s user=%s session=%s msg=%s",
        (request.mode or "fast"), (request.message_type or "TEXT"),
        (len(request.history) if request.history else 0),
        request.user_id, request.session_id,
        ((request.message or "")[:100] + ("…" if request.message and len(request.message) > 100 else "")),
    )
    # 注意：StreamingResponse 内部异步迭代时会在同一个 request 上下文里继续用同一个 req_id（FastAPI 请求上下文稳定）
    # 双 UUID 透传：
    #   - request.message_uuid → 传 stream_model_response 的 message_uuid（用户侧，更新用户正文/提取文本）
    #   - request.assistant_message_uuid → 传 assistant_message_uuid（assistant 侧，更新 searchResults/expertTrace）
    return StreamingResponse(
        stream_model_response(
            request.message,
            request.history,
            request.user_id,
            request.session_id,
            request.message_type or "TEXT",
            request.media_url,
            request.media_urls,
            request.message_uuid,           # 用户消息 UUID
            request.mode or "fast",
            request.assistant_message_uuid  # 新增：assistant 消息 UUID
        ),
        media_type="text/event-stream; charset=utf-8"
    )


# ==============================================================================
# 专家模式：编排器（单步计划 + analysis 思考历史）
# ==============================================================================

async def _expert_call_orchestrator_stream(
    state: ExpertState, result_holder: dict
) -> AsyncGenerator[str, None]:
    """
    流式调用专家编排器（V2）：
    - yield: analysis 的增量文本片段（供前端实时显示编排分析过程）
    - result_holder["result"]: 流结束后写入完整解析的 OrchResult 对象；失败时 result_holder["error"] 有值

    实现方式同快速模式：第一次 stream=True，重试改 stream=False
    """
    if not expert_orchestration_client:
        logger.warning("[专家模式][编排器] client 未初始化，兜底进入深度思考")
        result_holder["error"] = "编排器 client 未初始化"
        return

    # 构建对话历史（最近 10 条，供编排器参考上下文一致性）
    # 【硬编码移除】[:500] / 最近10条 逻辑 → 统一走 format_recent_history_for_llm（语义集中在 config.prompts.context）
    history_text = format_recent_history_for_llm(state.history)

    # 可用工具清单（动态注入，编排器 system prompt 里显式告知）
    # 【硬编码移除】原 for 循环里 web_search / memory_search 的 if/elif 分支 → default_tool_description_lines()
    tools_desc_lines = default_tool_description_lines(EXPERT_TOOL_REGISTRY)

    # 【硬编码移除】system_prompt / user_prompt → config.prompts.orchestrator_expert.build_*
    system_prompt = build_expert_orch_system_prompt(
        tools_description_lines=tools_desc_lines,
        max_iterations=state.max_iterations,
        current_iteration=state.iteration,
    )

    # 【硬编码移除】原 L1740-L1763 f-string 大段内联拼接 → build_expert_orch_user_prompt
    user_prompt = build_expert_orch_user_prompt(
        iteration_label=str(state.iteration + 1),
        max_iterations=state.max_iterations,
        thinking_history_text=_build_thinking_history_prompt(state),
        history_text=history_text,
        user_message=state.message,
        message_type=state.message_type,
        media_count=len(state.media_urls),
        extracted_text=state.extracted_text,
        memory_context_lines=state.memory_context_parts,
        search_context_parts=state.all_search_context_parts,
    )

    # 【硬编码移除】max_retries=2 → EXPERT_ORCH_MAX_RETRIES
    max_retries = EXPERT_ORCH_MAX_RETRIES
    # 第一次尝试用流式，后续重试改非流式（避免前端重复显示 analysis 文本）
    use_stream = True
    full_text = ""
    last_analysis_len = 0

    for attempt in range(1, max_retries + 1):
        try:
            response = await expert_orchestration_client.chat.completions.create(
                model=EXPERT_ORCHESTRATION_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                stream=use_stream,
                # 【硬编码移除】temperature=0.7 / max_tokens=3000 → EXPERT_ORCH_*
                temperature=EXPERT_ORCH_TEMPERATURE,
                max_tokens=EXPERT_ORCH_MAX_TOKENS,
                # 专家编排器输出 JSON，必须关思考：reasoning 会破坏 JSON 解析
                **_chat_kwargs_with_thinking(disable_thinking=True),
            )

            if use_stream:
                # ===== 流式模式：边接收边提取 analysis 增量 =====
                full_text = ""
                last_analysis_len = 0
                async for chunk in response:
                    if not chunk.choices or not chunk.choices[0].delta:
                        continue
                    delta = chunk.choices[0].delta.content
                    if not delta:
                        continue
                    full_text += delta
                    # 增量提取 analysis：只返回比上次多出来的部分
                    current = _try_extract_analysis(full_text, keys=("analysis",))
                    if len(current) > last_analysis_len:
                        yield current[last_analysis_len:]
                        last_analysis_len = len(current)
            else:
                # ===== 非流式重试模式 =====
                # Token 日志（只打日志，不存持久化，精简字段）
                if response.usage:
                    logger.info(
                        f"[编排器]Token消耗(模型:{EXPERT_ORCHESTRATION_MODEL}): "
                        f"prompt={response.usage.prompt_tokens}, completion={response.usage.completion_tokens}, "
                        f"total={response.usage.total_tokens}"
                    )
                if response.choices and response.choices[0].message and response.choices[0].message.content:
                    full_text = response.choices[0].message.content.strip()
                else:
                    full_text = ""

            # 解析完整 JSON
            cleaned = _clean_json_response(full_text)
            try:
                obj = json.loads(cleaned)
                result_holder["result"] = _parse_orch_result(obj, full_text)
                return
            except json.JSONDecodeError as je:
                logger.warning(f"[专家模式][编排器] 第{attempt}次 JSON 解析失败: {je}, 内容={cleaned[:120]}")
                if attempt < max_retries:
                    # 【硬编码移除】原 L1825-L1829 内联 3 行 → EXPERT_ORCH_USER_RETRY_APPEND
                    user_prompt += EXPERT_ORCH_USER_RETRY_APPEND
                    use_stream = False
                    continue
        except Exception as e:
            logger.error(f"[专家模式][编排器] 第{attempt}次调用失败({EXPERT_ORCHESTRATION_MODEL}): {e}")
            if attempt < max_retries:
                use_stream = False
                continue

    logger.error(f"[专家模式][编排器] 全部重试失败，兜底进入深度思考")
    result_holder["error"] = "编排器调用失败"


def _parse_orch_result(obj: dict, raw_output: str) -> OrchResult:
    """把编排器输出的 JSON 解析成 OrchResult 对象；非法值做规范化兜底"""
    action = obj.get("action") or EXPERT_ACTION_DEEP
    if action not in (EXPERT_ACTION_COLLECT, EXPERT_ACTION_DEEP, EXPERT_ACTION_REPLY):
        logger.warning(f"[专家模式][编排器] 非法 action={action}，兜底 deep_thinking")
        action = EXPERT_ACTION_DEEP

    step_obj = obj.get("step")
    plan_step: Optional[PlanStep] = None
    if action == EXPERT_ACTION_COLLECT:
        if isinstance(step_obj, dict) and step_obj.get("tools"):
            tools: List[ToolCall] = []
            for t in step_obj.get("tools", []):
                if isinstance(t, dict) and t.get("tool") and t.get("tool") in EXPERT_TOOL_REGISTRY:
                    tools.append(ToolCall(
                        tool=t["tool"],
                        params=dict(t.get("params") or {}),
                    ))
                else:
                    logger.warning(f"[专家模式][编排器] 忽略未知或非法工具项: {t}")
            if tools:
                plan_step = PlanStep(
                    purpose=str(step_obj.get("purpose") or "未描述目的"),
                    tools=tools,
                )
        if plan_step is None:
            # collect_tools 但 step 非法 → 放弃工具，强制深度思考（避免死循环）
            logger.warning("[专家模式][编排器] action=collect_tools 但没有可用工具，兜底 deep_thinking")
            action = EXPERT_ACTION_DEEP

    analysis = str(obj.get("analysis") or "")
    return OrchResult(action=action, step=plan_step, analysis=analysis, raw_output=raw_output)


# ==============================================================================
# 专家模式：EXECUTING 阶段（step.tools[] 并行工具执行）
# ==============================================================================

async def _expert_execute_step(state: ExpertState, step: PlanStep) -> Tuple[List[ToolExecutionRecord], List[str]]:
    """
    执行单步计划中的一组并行工具：
    - 用 asyncio.gather 并发；
    - 单个工具失败不影响其他工具；
    - 返回所有工具的执行记录（含 params/rawResult 不摘要）+ 前端流式消息行列表。
    """
    logger.info(f"[专家模式][EXECUTING] 本步 purpose={step.purpose}, 工具数={len(step.tools)}")
    coros: List[Awaitable[Tuple[ToolExecutionRecord, List[str]]]] = []
    for tc in step.tools:
        if tc.tool not in EXPERT_TOOL_REGISTRY:
            # 未知工具：构造失败记录，不抛异常，保证编排器下轮能感知
            now_ms = 0
            failed = ToolExecutionRecord(
                tool=tc.tool,
                params=tc.params,
                success=False,
                durationMs=now_ms,
                resultCount=None,
                error=f"未知工具: {tc.tool}",
                rawResult=None,
            )
            coros.append(asyncio.sleep(0, result=(failed, [])))  # 用 0 秒休眠包成 Awaitable
            continue
        fn, _desc = EXPERT_TOOL_REGISTRY[tc.tool]
        coros.append(fn(tc.params, state))

    results: List[Tuple[ToolExecutionRecord, List[str]]] = await asyncio.gather(*coros, return_exceptions=False)
    records: List[ToolExecutionRecord] = []
    stream_lines: List[str] = []
    for rec, lines in results:
        records.append(rec)
        stream_lines.extend(lines)
    # 工具执行概况：汇总成一行（工具多时避免多条线性刷屏）；详情要排查时再开 DEBUG 或看 ExpertState.orch_history
    tool_quick = ", ".join(
        f"{r.tool}:{'OK' if r.success else 'FAIL'},{r.durationMs}ms"
        for r in records
    )
    logger.info(
        "[专家模式][EXECUTING] 完成 purpose=%s tools=[%s] total_ms=%.0f",
        step.purpose, tool_quick, sum(max(r.durationMs or 0, 0) for r in records),
    )
    return records, stream_lines


# ==============================================================================
# 专家模式：FINAL 阶段两条流式路径（chunk 到就立即 yield，零等待）
# ==============================================================================

async def _expert_final_deep_thinking(state: ExpertState) -> AsyncIterator[str]:
    """
    FINAL: deep_thinking 流式路径（硬校验：最多调用一次）。
    每个 chunk 立即 yield 两条：type=thinking（前端黄色面板）和 type=content（最终正文）。
    深度思考输出 = 最终回复正文，不再二次润色。

    【硬编码移除】
      - state.final_path 字面量 → EXPERT_FINAL_PATH_DEEP
      - system_prompt → DEEP_THINKING_SYSTEM_PROMPT
      - thinking_start.message → DEEP_THINKING_START_MESSAGE
      - temperature → DEEP_THINKING_TEMPERATURE
      - reasoning_effort → DEEP_THINKING_REASONING_EFFORT
    """
    assert not state.deep_thinking_called, "深度思考重复调用（状态机硬校验失败）"
    state.deep_thinking_called = True
    # 【硬编码移除】"deep_thinking" → EXPERT_FINAL_PATH_DEEP
    state.final_path = EXPERT_FINAL_PATH_DEEP
    logger.info("[专家模式][FINAL] 进入深度思考流式路径")

    if not deep_thinking_client:
        logger.warning("[专家模式][深度思考] client 未初始化，降级为简单错误提示")
        fallback = "（深度思考模型未初始化，请检查配置）"
        state.assistant_response = fallback
        yield _json_line({"type": "thinking_error", "data": {"error": fallback}})
        yield _json_line({"type": "content", "data": fallback})
        return

    start_time = time.time()
    global_context = _build_global_context(state)

    # 构造多模态消息（如需要）
    user_content: List[Dict[str, Any]] = [{"type": "text", "text": global_context}]
    for url in state.media_urls:
        user_content.append({"type": "image_url", "image_url": {"url": url}})

    # 【硬编码移除】深度思考 system_prompt → DEEP_THINKING_SYSTEM_PROMPT
    system_prompt = DEEP_THINKING_SYSTEM_PROMPT

    # 思考开始提示块
    # 【硬编码移除】"进入深度思考，正在逐步分析..." → DEEP_THINKING_START_MESSAGE
    yield _json_line({"type": "thinking_start", "data": {
        "message": DEEP_THINKING_START_MESSAGE
    }})

    chunk_count = 0
    try:
        response = await deep_thinking_client.chat.completions.create(
            model=EXPERT_DEEP_THINKING_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            stream=True,
            # 【硬编码移除】temperature=0.3 → DEEP_THINKING_TEMPERATURE
            temperature=DEEP_THINKING_TEMPERATURE,
            # 【硬编码移除】reasoning_effort="low" → DEEP_THINKING_REASONING_EFFORT
            **_chat_kwargs_with_thinking(disable_thinking=False,
                                         reasoning_effort=DEEP_THINKING_REASONING_EFFORT),
        )
        logger.info("[深度思考] 流式连接建立 model=%s", EXPERT_DEEP_THINKING_MODEL)
        async for chunk in response:
            if chunk.choices and chunk.choices[0].delta:
                delta = chunk.choices[0].delta
                # 优先从 reasoning_content / thinking_content 读取"思考栏"文本：
                # 深度思考模型（qwen3.8-27b）会把推理链放到 reasoning_content 字段，
                # 正式回复正文放到 content 字段；旧端点如果没有 reasoning 字段，
                # 会把推理链和正文一起塞进 content，此时再退化为 content→thinking+content 双发
                reasoning = (
                    getattr(delta, "reasoning_content", None)
                    or getattr(delta, "thinking_content", None)
                    or getattr(delta, "reasoning_delta", None)
                )
                content = delta.content
                if reasoning:
                    # 累积深度思考推理链全文，用于持久化到 expertTrace
                    state.deep_thinking_reasoning += reasoning
                    yield _json_line({"type": "thinking", "data": reasoning})
                if content:
                    chunk_count += 1
                    state.assistant_response += content
                    yield _json_line({"type": "content", "data": content})
                    # 修复：不再把 content 双发到 thinking 频道
                    # 之前旧逻辑"无 reasoning 时把 content 也追加到思考栏"会导致
                    # 前端思考过程和最终回复完全重复，用户反馈此问题
        duration = time.time() - start_time
        if hasattr(response, "usage") and response.usage:
            logger.info(
                "[深度思考] Token消耗 model=%s prompt=%s completion=%s total=%s dur=%.2fs",
                EXPERT_DEEP_THINKING_MODEL,
                response.usage.prompt_tokens, response.usage.completion_tokens,
                response.usage.total_tokens, duration,
            )
        # 深度思考：一行聚合（替代原来多行详情）
        logger.info(
            "[深度思考] 完成 chunks=%s reply_len=%s dur=%.2fs",
            chunk_count, len(state.assistant_response), duration,
        )
    except Exception as e:
        duration = time.time() - start_time
        logger.error(f"[专家模式][深度思考] 失败: {e}, 耗时={duration:.2f}秒")
        yield _json_line({"type": "thinking_error", "data": {"error": str(e)}})
        if not state.assistant_response:
            fallback = f"（深度思考过程出错：{str(e)[:200]}）"
            state.assistant_response = fallback
            yield _json_line({"type": "content", "data": fallback})


async def _expert_final_reply_model(state: ExpertState) -> AsyncIterator[str]:
    """
    FINAL: reply_directly 路径 —— 直接调用编排器模型（或回复模型）流式生成正文。
    简单问题专用（问候、极短常识），不经过深度思考。
    流式 chunk 到即 yield type=content，零等待。

    【硬编码移除】
      - state.final_path 字面量 → EXPERT_FINAL_PATH_REPLY
      - system_prompt → REPLY_DIRECTLY_SYSTEM_PROMPT
      - user_prompt → build_reply_directly_user_prompt
      - temperature / max_tokens → REPLY_DIRECTLY_*
    """
    # 【硬编码移除】"reply_directly" → EXPERT_FINAL_PATH_REPLY
    state.final_path = EXPERT_FINAL_PATH_REPLY
    logger.info("[专家模式][FINAL] 进入简单回复模型流式路径")

    if not expert_orchestration_client:
        fallback = "（回复模型未初始化）"
        state.assistant_response = fallback
        yield _json_line({"type": "content", "data": fallback})
        return

    start_time = time.time()
    global_context = _build_global_context(state)
    # 【硬编码移除】[:500] / 最近10条 内联循环 → format_recent_history_for_llm
    history_text = format_recent_history_for_llm(state.history)

    # 【硬编码移除】原 L1992-L1997 内联 system prompt → REPLY_DIRECTLY_SYSTEM_PROMPT
    system_prompt = REPLY_DIRECTLY_SYSTEM_PROMPT
    # 【硬编码移除】原 L1998-L2003 内联 user prompt f-string → build_reply_directly_user_prompt
    user_prompt = build_reply_directly_user_prompt(
        user_message=state.message,
        global_context=global_context,
        history_text=history_text,
    )

    try:
        response = await expert_orchestration_client.chat.completions.create(
            model=EXPERT_ORCHESTRATION_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            stream=True,
            # 【硬编码移除】temperature=0.7 / max_tokens=2000 → REPLY_DIRECTLY_*
            temperature=REPLY_DIRECTLY_TEMPERATURE,
            max_tokens=REPLY_DIRECTLY_MAX_TOKENS,
            # 用户要求：仅深度思考模型开思考；reply_directly（简单问题）属于快速回复路径，关思考
            # reasoning_effort 这里保留 "low" 常量语义：config.behavior 未单独拆这一项（影响极小）
            **_chat_kwargs_with_thinking(disable_thinking=True, reasoning_effort="low"),
        )
        chunk_count = 0
        async for chunk in response:
            if chunk.choices and chunk.choices[0].delta:
                delta = chunk.choices[0].delta
                reasoning = (
                    getattr(delta, "reasoning_content", None)
                    or getattr(delta, "thinking_content", None)
                    or getattr(delta, "reasoning_delta", None)
                )
                piece = delta.content
                if reasoning:
                    yield _json_line({"type": "thinking", "data": reasoning})
                if piece:
                    chunk_count += 1
                    state.assistant_response += piece
                    yield _json_line({"type": "content", "data": piece})
        duration = time.time() - start_time
        if hasattr(response, "usage") and response.usage:
            logger.info(
                "[回复模型] Token消耗 model=%s prompt=%s completion=%s total=%s dur=%.2fs",
                EXPERT_ORCHESTRATION_MODEL,
                response.usage.prompt_tokens, response.usage.completion_tokens,
                response.usage.total_tokens, duration,
            )
        logger.info(
            "[回复模型] 完成 chunks=%s reply_len=%s dur=%.2fs",
            chunk_count, len(state.assistant_response), duration,
        )
    except Exception as e:
        logger.error(f"[专家模式][回复模型] 失败: {e}")
        fallback = f"（简单回复模型出错：{str(e)[:200]}）"
        state.assistant_response = fallback
        yield _json_line({"type": "content", "data": fallback})


# ==============================================================================
# 专家模式：持久化（等回复完成后，一次性回写 expertTrace 到后端）
# ==============================================================================

def _build_expert_trace_payload(state: ExpertState) -> Dict[str, Any]:
    """
    构造 expertTrace 对象：
    - 工具调用结果不做摘要，保留 params / rawResult；
    - 编排器每次的 analysis 都保存（你下轮编排器看的思考历史就来源于此）。
    - 不保存 orch tokens / orch duration（只在日志里打）—— 精简无用字段。
    字段名都是完整语义名，不做缩写。
    """
    history_arr: List[Dict[str, Any]] = []
    for rec in state.orch_history:
        history_arr.append({
            "iteration": rec.iteration,
            "action": rec.action,
            "purpose": rec.purpose,
            "analysis": rec.analysis,
            "tools": [
                {
                    "tool": t.tool,
                    "params": t.params,                 # 不摘要，保留原始入参
                    "success": t.success,
                    "durationMs": t.durationMs,
                    "resultCount": t.resultCount,
                    "error": t.error,
                    "rawResult": t.rawResult,           # 不摘要，保留精简原始结果
                }
                for t in rec.tools
            ],
        })
    payload = {
        "finalPath": state.final_path,
        "iterationCount": len([1 for r in state.orch_history if r.action]),
        "history": history_arr,
        # 深度思考推理链全文（与工具调用结果同等持久化，用户要求）
        "deepThinkingReasoning": state.deep_thinking_reasoning,
    }
    return payload


async def update_backend_expert_trace(session_id: Optional[int], message_uuid: Optional[str], expert_trace: Dict[str, Any]) -> bool:
    """
    专家模式 FINAL 结束后调用：把 expertTrace 回写进 assistant 消息 JSON（Redis + MySQL 会话详情表）。

    实现策略说明（解决 401 的核心）：
    👉 不再使用新增的 /api/message/update-expert-trace 独立 URL（需要同步改 InterceptorConfig 白名单 + 重新编译 jar，
       用户当前运行的旧 jar 没有白名单就会反复 401）。
    👉 改为复用现成白名单的 /api/message/update-content URL：Token 拦截器 exclude + InternalApiInterceptor
       addPathPatterns 两处都早已注册，从来不会 401（这就是你说的"现成回写函数"）。
    参数：只传 expert_trace（JSON 字符串化），message="" 与 extracted_text="" 由后端 "非空才覆盖" 策略忽略。
    """
    if not message_uuid or not session_id:
        logger.warning("[专家模式][持久化] session_id 或 message_uuid 为空，跳过 expertTrace 回写")
        return False
    try:
        payload = {
            "session_id": str(session_id),
            "message_uuid": message_uuid,
            "message": "",                    # 后端：空字符串 → 不覆盖原 content
            "extracted_text": "",              # 后端：空字符串 → 不覆盖原 extractedText
            "expert_trace": json.dumps(expert_trace, ensure_ascii=False),
        }
        # 统一格式：X-Internal-Secret + Content-Type
        headers = {
            "Content-Type": "application/json",
            "X-Internal-Secret": INTERNAL_SECRET,
        }
        # 【硬编码移除】"/api/message/update-content" → BACKEND_ENDPOINT_UPDATE_MESSAGE_CONTENT
        update_url = f"{BACKEND_BASE_URL}{BACKEND_ENDPOINT_UPDATE_MESSAGE_CONTENT}"
        logger.debug(
            "[专家模式][持久化] 复用现成 update-content 通路回写 expertTrace: url=%s, sessionId=%s, messageUuid=%s, payloadLen=%s",
            update_url, session_id, message_uuid, len(payload["expert_trace"]),
        )
        # 【硬编码移除】Timeout(10.0, connect=3.0) → BACKEND_INTERNAL_API_TIMEOUT + BACKEND_INTERNAL_API_CONNECT_TIMEOUT
        async with httpx.AsyncClient(timeout=httpx.Timeout(
            BACKEND_INTERNAL_API_TIMEOUT,
            connect=BACKEND_INTERNAL_API_CONNECT_TIMEOUT,
        )) as client:
            resp = await client.post(update_url, json=payload, headers=headers)
            # update-content 返回的 JSON: {code,message}；与 update-expert-trace 之前使用的"code==200且HTTP 200"保持一致
            ok = resp.status_code == 200
            if ok:
                try:
                    body = resp.json()
                    ok = (body.get("code") == 200)
                except Exception:
                    # 兜底：非 JSON 情况下 HTTP 200 算成功
                    pass
            if ok:
                logger.info(
                    "[专家模式][持久化] expertTrace 回写成功(复用update-content通路), "
                    "history_length=%s, finalPath=%s",
                    len(expert_trace.get("history", [])), expert_trace.get("finalPath"),
                )
            else:
                logger.error(
                    "[专家模式][持久化] expertTrace 回写失败(复用update-content通路): status=%s, body=%s",
                    resp.status_code, resp.text[:200],
                )
            return ok
    except Exception as e:
        logger.error("[专家模式][持久化] expertTrace 回写异常: %s: %s", type(e).__name__, e)
        return False


# ==============================================================================
# 专家模式主入口：状态机驱动的处理流程（替代原 expert_mode_process）
# ==============================================================================

async def expert_mode_process(message: str, history: Optional[List[dict]], user_id: Optional[int],
                              session_id: Optional[int], message_type: str, media_url: Optional[str],
                              media_urls: Optional[List[str]], message_uuid: Optional[str],
                              assistant_message_uuid: Optional[str],
                              request_start_time: float):
    """
    专家模式处理流程（状态机 V2）：
      预处理（媒体提取+记忆检索并行）
        → DECIDING ↔ EXECUTING 循环（每次编排器输出单步计划；每步工具并行）
        → FINAL（deep_thinking 流式 或 reply_model 流式）
        → DONE（一次性持久化 expertTrace 到 assistant 消息 JSON）
    硬约束：
      - 深度思考最多一次，且只能 FINAL 阶段调用；
      - 编排器一次只生成一步，之后靠自己写的 analysis 回忆之前思路；
      - 编排器不输出回复正文，回复只在 FINAL 流式生成。
    参数说明：
      - message_uuid:           用户消息 UUID（role=user）→ 更新用户侧字段（正文/提取文本）
      - assistant_message_uuid: assistant 回复 UUID（role=assistant）→ 写入 expertTrace/searchResults
    """
    logger.info("[专家模式] 开始处理 userUuid=%s assistantUuid=%s",
        (message_uuid or "")[:8] + "…" if message_uuid else "-",
        (assistant_message_uuid or "")[:8] + "…" if assistant_message_uuid else "-",
    )

    # ========== 1. 预处理：媒体提取 + 记忆检索 并行 ==========
    all_media_urls: List[str] = list(media_urls or [])
    if media_url and media_url not in all_media_urls:
        all_media_urls.append(media_url)

    # 构造 state（状态机唯一数据源）
    # 【关键语义调整】ExpertState.message_uuid 现在是 assistant_message_uuid（assistant 侧字段回写用）
    # 新增 user_message_uuid 保存用户消息 UUID（更新用户内容/提取文本时用）
    # 兼容：旧 jar 包（还没升级双 UUID 改造）不会传 assistant_message_uuid，
    #       fallback 到用户 UUID，并打印 WARN 提示升级 jar（否则回写可能被 Service 校验 role 拒绝）。
    final_assistant_uuid = assistant_message_uuid
    if not final_assistant_uuid:
        logger.warning("[专家模式] 后端未传 assistant_message_uuid，fallback 到 userUuid=%s。"
            "请升级后端 jar（双 UUID 版本），否则 expertTrace 回写可能因 role 校验失败被拒绝",
            (message_uuid or "")[:8] + "…" if message_uuid else "-",
        )
        final_assistant_uuid = message_uuid
    state = ExpertState(
        message=message,
        message_type=message_type,
        media_urls=all_media_urls,
        history=list(history or []),
        user_id=user_id,
        session_id=str(session_id) if session_id else None,
        user_message_uuid=message_uuid,                # 用户消息 UUID
        message_uuid=final_assistant_uuid,             # 原字段语义：assistant UUID（用于 expertTrace 回写）
        request_start_time=request_start_time,
        max_iterations=EXPERT_MAX_ITERATIONS,
    )

    # 媒体提取（图片/文件）
    is_multimodal = message_type in ("IMAGE", "FILE")
    extract_task: Optional[Awaitable[str]] = None
    if is_multimodal and all_media_urls:
        if len(all_media_urls) > 1:
            extract_task = extract_media_text_batch(all_media_urls, message_type)
        else:
            extract_task = extract_media_text(all_media_urls[0], message_type)

    # 记忆检索（用原始提问文本，并行发起）
    mem_task: Optional[Awaitable[List[dict]]] = None
    if session_id:
        async def _mem_wrap() -> List[dict]:
            nonlocal state
            t0 = time.time()
            try:
                q = message if not state.extracted_text else f"{message} {state.extracted_text}"
                items = await memory_service.search_memories(q, str(session_id), top_k=5)
                logger.info(f"[专家模式][记忆检索] 完成，{len(items)}条, 耗时={time.time()-t0:.2f}秒")
                return items
            except Exception as e:
                logger.error(f"[专家模式][记忆检索] 失败: {e}")
                return []
        mem_task = _mem_wrap()

    # 并行等：媒体提取 + 记忆检索
    if extract_task and mem_task:
        (state.extracted_text, state.memories) = await asyncio.gather(extract_task, mem_task)
    elif extract_task:
        state.extracted_text = await extract_task
    elif mem_task:
        state.memories = await mem_task

    if state.extracted_text:
        logger.info(f"[专家模式][媒体提取] 成功，长度={len(state.extracted_text)}字符")
        # 注意：如果要把 extracted_text 回写进用户消息 JSON（update_backend_message_content），
        # 必须使用 state.user_message_uuid（用户消息的 UUID，对应 role=user），
        # 不要用 state.message_uuid（现在语义是 assistant 消息 UUID，否则会被后端 Service 校验 role 拒绝）。
    if state.memories:
        # 把 INIT 阶段的基础记忆塞到 memory_context_parts 里给后续编排器/FINAL 用
        state.memory_context_parts.extend([f"- {m.get('content','')[:300]}" for m in state.memories])

    # ========== 2. 状态机主循环：DECIDING ↔ EXECUTING ==========
    # 最大 collect_tools 迭代次数限制（超过强制 FINAL deep_thinking）
    while state.phase in (ExpertPhase.DECIDING, ExpertPhase.EXECUTING):

        # -------------------- DECIDING：编排器决策单步 --------------------
        if state.phase == ExpertPhase.DECIDING:
            state.iteration += 1
            logger.info(
                f"[专家模式][DECIDING] 第 {state.iteration}/{state.max_iterations} 次编排"
            )
            # 流式调用编排器：边生成边 yield analysis 增量给前端，实现"打字机"效果
            orch_result_holder = {}
            async for analysis_delta in _expert_call_orchestrator_stream(state, orch_result_holder):
                yield _json_line({"type": "orchestration_chunk", "data": {"delta": analysis_delta}})

            orch = orch_result_holder.get("result") or OrchResult(
                action=EXPERT_ACTION_DEEP, step=None,
                analysis="编排器调用失败，兜底进入深度思考", raw_output="",
            )
            state.pending_orch = orch
            logger.info(
                f"[专家模式][DECIDING] 编排结果: action={orch.action}, "
                f"step_purpose={(orch.step.purpose if orch.step else None)}, "
                f"tools={[t.tool for t in (orch.step.tools if orch.step else [])]}, "
                f"analysis={orch.analysis[:80]}"
            )

            # ===== 编排步骤完成：通知前端结构化信息（analysis 已通过 chunk 流式显示）=====
            yield _json_line({"type": "orchestration_step", "data": {
                "iteration": state.iteration,
                "phase": "planning" if orch.action == EXPERT_ACTION_COLLECT else "thinking",
                "action": orch.action,
                "purpose": (orch.step.purpose if orch.step else None),
                # analysis 不再重复发送（已通过 orchestration_chunk 逐步显示给前端）
            }})

            # 本轮 orch 先填一个占位 history record，等 EXECUTING（或 FINAL）结束后再回填 tools 执行结果
            rec = OrchHistoryRecord(
                iteration=state.iteration,
                action=orch.action,
                purpose=(orch.step.purpose if orch.step else None),
                analysis=orch.analysis,
                tools=[],
            )
            state.orch_history.append(rec)

            # 超过最大 collect 迭代次数 → 强制跳 FINAL(deep)
            if orch.action == EXPERT_ACTION_COLLECT and state.iteration >= state.max_iterations:
                logger.warning(
                    f"[专家模式][硬校验] 达到最大迭代次数 {state.max_iterations}，"
                    f"跳过 collect_tools，强制 FINAL(deep_thinking)"
                )
                orch.action = EXPERT_ACTION_DEEP
                state.pending_orch = orch
                # 修改刚塞进去的 history record 的 action
                state.orch_history[-1].action = EXPERT_ACTION_DEEP
                state.orch_history[-1].purpose = "（系统强制进入深度思考）"

            # 转移到对应阶段
            if orch.action == EXPERT_ACTION_COLLECT:
                state.pending_step = orch.step
                state.phase = ExpertPhase.EXECUTING
                continue
            elif orch.action == EXPERT_ACTION_DEEP:
                state.phase = ExpertPhase.FINAL
                # 【硬编码移除】"deep_thinking" → EXPERT_FINAL_PATH_DEEP
                state.final_path = EXPERT_FINAL_PATH_DEEP
                break  # 跳出主循环到 FINAL
            elif orch.action == EXPERT_ACTION_REPLY:
                state.phase = ExpertPhase.FINAL
                # 【硬编码移除】"reply_directly" → EXPERT_FINAL_PATH_REPLY
                state.final_path = EXPERT_FINAL_PATH_REPLY
                break  # 跳出主循环到 FINAL
            else:
                # 理论上 _parse_orch_result 已经兜底了，这里保险
                logger.warning(f"[专家模式][硬校验] 未知 action={orch.action}，兜底 FINAL(deep)")
                state.phase = ExpertPhase.FINAL
                # 【硬编码移除】"deep_thinking" → EXPERT_FINAL_PATH_DEEP
                state.final_path = EXPERT_FINAL_PATH_DEEP
                break

        # -------------------- EXECUTING：并行跑一组工具，立刻回到 DECIDING --------------------
        if state.phase == ExpertPhase.EXECUTING:
            step = state.pending_step
            if step is None:
                logger.warning("[专家模式][EXECUTING] pending_step 为空，退回 DECIDING")
                state.phase = ExpertPhase.DECIDING
                continue

            tool_records, stream_lines = await _expert_execute_step(state, step)

            # 先向前端 flush 流式状态块（search_start / search_summary ...）
            for line in stream_lines:
                yield line + ("" if line.endswith("\n") else "\n")

            # 回填到本 iteration 的 OrchHistoryRecord.tools（DECIDING 阶段塞了一个空 tools[] 的占位 record）
            state.orch_history[-1].tools = tool_records

            # 每步执行完立刻回到 DECIDING，让编排器重新分析（可能继续下一步或直接进入 FINAL）
            state.phase = ExpertPhase.DECIDING
            continue

    # ========== 3. FINAL：两条流式路径之一 ==========
    logger.info(f"[专家模式] 进入 FINAL 阶段，final_path={state.final_path}")
    # 【硬编码移除】字面量 "deep_thinking" → EXPERT_FINAL_PATH_DEEP
    if state.final_path == EXPERT_FINAL_PATH_DEEP:
        async for line in _expert_final_deep_thinking(state):
            yield line
    else:
        async for line in _expert_final_reply_model(state):
            yield line

    # ========== 4. 通过 SSE 流发送 expertTrace（替代 HTTP 回调，避免竞态） ==========
    # 后端在 SSE 流中捕获 expert_trace 事件，流结束时随 saveMessage 一起持久化
    # 不再使用 asyncio.create_task(update_backend_expert_trace) 避免 Redis 扫不到消息的竞态
    expert_trace_payload = _build_expert_trace_payload(state)
    yield _json_line({
        "type": "expert_trace",
        "data": json.dumps(expert_trace_payload, ensure_ascii=False),
    })

    # ========== 6. 异步保存记忆（保持原逻辑：用户消息 + 提取内容 + 最终回复） ==========
    if user_id and state.assistant_response:
        if state.extracted_text:
            combined_content = (
                f"[用户提问]{message}\n[图片内容]{state.extracted_text}"
                if message_type == "IMAGE" else
                f"[用户提问]{message}\n[文件内容]{state.extracted_text}"
            )
            asyncio.create_task(memory_service.async_extract_and_store(
                combined_content, state.assistant_response, user_id,
                int(state.session_id) if state.session_id and str(state.session_id).isdigit() else None,
            ))
            logger.info(
                f"[专家模式][记忆持久化] 启动异步任务（含媒体提取）, "
                f"消息长度={len(message)}, 提取长度={len(state.extracted_text)}, 回复长度={len(state.assistant_response)}"
            )
            if state.user_message_uuid and state.session_id:
                try:
                    # 修复：必须用 user_message_uuid（用户消息UUID），
                    # 不能用 message_uuid（现在是assistant UUID，会被role校验拒绝）
                    await update_backend_message_content(
                        int(state.session_id) if str(state.session_id).isdigit() else None,
                        state.user_message_uuid, message, state.extracted_text,
                    )
                except Exception as e:
                    logger.warning(f"[专家模式] 回写图片提取内容失败: {e}")
        else:
            asyncio.create_task(memory_service.async_extract_and_store(
                message, state.assistant_response, user_id,
                int(state.session_id) if state.session_id and str(state.session_id).isdigit() else None,
            ))
            logger.info(
                f"[专家模式][记忆持久化] 启动异步任务（无媒体提取）, "
                f"消息长度={len(message)}, 回复长度={len(state.assistant_response)}"
            )

    # ========== 7. 汇总日志（一行紧凑，替代原来 8 行分隔线 + 多行分类）==========
    total_duration = time.time() - request_start_time
    logger.info(
        "[专家模式][请求结束] dur=%.2fs path=%s reply=%s search=%s orch=%s tools=%s session=%s",
        total_duration, state.final_path,
        len(state.assistant_response),
        len(state.search_results_flattened),
        state.iteration,
        sum(len(r.tools) for r in state.orch_history),
        state.session_id,
    )
    state.phase = ExpertPhase.DONE


async def generate_summary(messages: Optional[List[dict]], existing_summary: Optional[str]):
    """
    对话摘要生成：把 messages（含 IMAGE/FILE 媒体类型、extractedText）拼接成 LLM 可读文本，
    再走结构化摘要模型。输出 JSON：{key_points, entities, summary}。

    【硬编码移除】
      - 媒体消息的 prefix（图片/文件上传的 4 种标题行）→ HISTORY_IMAGE_* / HISTORY_FILE_*
      - system_prompt / user_prompt → config.prompts.summarize SUMMARIZE_SYSTEM_PROMPT / build_summary_user_prompt
      - temperature → SUMMARIZE_TEMPERATURE；关思考 → SUMMARIZE_DISABLE_THINKING
    """
    logger.info(f"[摘要生成] 开始: 消息数量={len(messages) if messages else 0}, 已有摘要={'是' if existing_summary else '否'}")

    messages_text = ""
    if messages:
        for msg in messages:
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            msg_type = msg.get("messageType", "TEXT")
            extracted_text_history = msg.get("extractedText")

            # 【硬编码移除】4 段内联 f-string（"[用户上传了图片]..."）→ 统一走 config.prompts.context 常量
            if msg_type == "IMAGE":
                if extracted_text_history:
                    content = HISTORY_IMAGE_PREFIX_WITH_EXTRACTED.format(
                        content=content, extracted_text=extracted_text_history
                    )
                else:
                    content = HISTORY_IMAGE_PREFIX_SIMPLE.format(content=content)
            elif msg_type == "FILE":
                if extracted_text_history:
                    content = HISTORY_FILE_PREFIX_WITH_EXTRACTED.format(
                        content=content, extracted_text=extracted_text_history
                    )
                else:
                    content = HISTORY_FILE_PREFIX_SIMPLE.format(content=content)
            messages_text += f"{role}: {content}\n"

    # 【硬编码移除】原 L2458-L2465 大段 system prompt → SUMMARIZE_SYSTEM_PROMPT
    system_prompt = SUMMARIZE_SYSTEM_PROMPT

    # 【硬编码移除】原 L2467-L2475 内联 user prompt f-string → build_summary_user_prompt
    user_prompt = build_summary_user_prompt(
        messages_text=messages_text,
        existing_summary=existing_summary,
    )

    try:
        response = await client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            stream=False,
            # 【硬编码移除】temperature=0.3 → SUMMARIZE_TEMPERATURE
            temperature=SUMMARIZE_TEMPERATURE,
            # 摘要生成输出 JSON，必须关思考：避免 reasoning 混入摘要文本
            **_chat_kwargs_with_thinking(disable_thinking=SUMMARIZE_DISABLE_THINKING),
        )

        if response.choices and response.choices[0].message and response.choices[0].message.content:
            summary_content = response.choices[0].message.content.strip()
            logger.info(f"[摘要生成] 成功: 长度={len(summary_content)}")
            return summary_content
        else:
            logger.warning("[摘要生成] 未收到模型返回的摘要内容")
            return None

    except Exception as e:
        logger.error(f"[摘要生成] 出错: {str(e)}")
        return None


@app.post("/summarize")
async def summarize(request: SummarizeRequest):
    logger.info(f"[API入口] 收到摘要请求: 消息数量={len(request.messages) if request.messages else 0} 条")

    summary = await generate_summary(request.messages, request.existing_summary)

    if summary:
        return {"summary": summary}
    else:
        return {"summary": ""}


from fastapi import HTTPException, Request
from qdrant_client import models as qdrant_models
from config.settings import QDRANT_COLLECTION

class DeleteMemoryRequest(BaseModel):
    session_id: str

@app.delete("/memory/delete")
async def delete_memory(request: DeleteMemoryRequest, request_obj: Request):
    # 统一使用顶部全局 INTERNAL_SECRET，避免散落 os.getenv 导致 NameError
    header_secret = request_obj.headers.get("X-Internal-Secret", "")

    if not INTERNAL_SECRET or not header_secret or not header_secret == INTERNAL_SECRET:
        logger.warning(f"[内存删除] 内部接口认证失败: sessionId={request.session_id}")
        raise HTTPException(status_code=403, detail="内部接口认证失败")
    
    logger.info(f"[内存删除] 收到请求: sessionId={request.session_id}, 类型={type(request.session_id)}")
    
    try:
        if not memory_service._initialized:
            memory_service.init()
        
        qdrant_client = memory_service.qdrant_client
        
        result = qdrant_client.delete(
            collection_name=QDRANT_COLLECTION,
            points_selector=qdrant_models.FilterSelector(
                filter=qdrant_models.Filter(
                    must=[qdrant_models.FieldCondition(
                        key="sessionId",
                        match=qdrant_models.MatchValue(value=str(request.session_id))
                    )]
                )
            )
        )
        
        logger.info(f"[内存删除] 完成: sessionId={request.session_id}, 已删除{result.count if hasattr(result, 'count') else '未知'}条记录")
        
        return {"success": True, "message": "记忆删除成功"}
    except Exception as e:
        logger.error(f"[内存删除] 失败: sessionId={request.session_id}, 错误={str(e)}")
        raise HTTPException(status_code=500, detail=f"删除记忆失败: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    logger.info("启动XiaoAi Agent...")
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)